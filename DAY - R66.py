#!/usr/bin/env python3
# -*- coding: utf-8 -*-
# Project : DAY - R 
# GLOBAL CONFLICTS & THREATS DETECTION & ANALYSIS INTELLIGENCE SYSTEM (GCTDAI - System) 
# Department of Defense - Joint Intelligence Command
# TS//SCI Clearance Required - Codeword: BLACKSTAR


import os
import sys
import subprocess
import re
import asyncio
import ssl
import datetime
import platform
import json
import hashlib
import threading
import time
import warnings
import tempfile
import webbrowser
import shutil
import zipfile
import socket
import uuid
import base64
import importlib
import traceback
from urllib.parse import urlparse, urljoin
from collections import defaultdict, deque
import random
import stat
import logging
from logging.handlers import RotatingFileHandler
import configparser
import queue

# --- Enhanced Environment Setup ---
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
VENV_DIR = os.path.join(SCRIPT_DIR, "mgcis_venv")
REPORT_DIR = os.path.join(SCRIPT_DIR, "reports")
DB_DIR = os.path.join(SCRIPT_DIR, "database")
HTML_DIR = os.path.join(SCRIPT_DIR, "dashboards")
TOOL_DIR = os.path.join(SCRIPT_DIR, "custom_tools")
CONFIG_DIR = os.path.join(SCRIPT_DIR, "config")
LOG_DIR = os.path.join(SCRIPT_DIR, "logs")
AUDIO_DIR = os.path.join(SCRIPT_DIR, "audio_briefings")
DASHBOARD_DIR = os.path.join(SCRIPT_DIR, "live_dashboards")

# --- Fixed Indentation Fixer ---
def auto_fix_indentation(code):
    """Automatically fix common indentation errors in Python code"""
    lines = code.splitlines()
    fixed_lines = []
    indent_level = 0
    in_block = False
    indent_size = 4
    
    for line in lines:
        stripped = line.strip()
        leading_spaces = len(line) - len(line.lstrip())
        
        # Handle dedent keywords
        if stripped in ['return', 'break', 'continue', 'pass']:
            indent_level = max(0, indent_level - 1)
            in_block = False
            
        # Calculate expected indent
        expected_indent = indent_level * indent_size
        
        # Fix current line's indentation
        if leading_spaces != expected_indent:
            fixed_line = (' ' * expected_indent) + stripped
        else:
            fixed_line = line
            
        # Handle block starts
        if stripped.endswith(':') and not stripped.startswith(('"""', "'''")):
            in_block = True
            indent_level += 1
            
        # Handle block endings
        if stripped and not in_block and indent_level > 0:
            indent_level = max(0, indent_level - 1)
            
        fixed_lines.append(fixed_line)
        
    return '\n'.join(fixed_lines)

# --- Custom JSON Encoder for datetime objects ---
class DateTimeEncoder(json.JSONEncoder):
    """Custom JSON encoder that handles datetime objects"""
    def default(self, obj):
        if isinstance(obj, datetime.datetime):
            return obj.isoformat()
        return json.JSONEncoder.default(self, obj)

# Create directories with secure permissions (0700)
def create_secure_directory(path):
    """Create directory with secure permissions if it doesn't exist"""
    if not os.path.exists(path):
        os.makedirs(path, mode=0o700, exist_ok=True)
        print(f"🔒 Created secure directory: {path}")
    else:
        # Fix permissions if directory exists
        os.chmod(path, 0o700)
        print(f"🔒 Fixed permissions for: {path}")

for directory in [VENV_DIR, REPORT_DIR, DB_DIR, HTML_DIR, TOOL_DIR, CONFIG_DIR, LOG_DIR, AUDIO_DIR, DASHBOARD_DIR]:
    create_secure_directory(directory)

# --- Virtual Environment Management ---
def setup_virtual_environment():
    """Create and activate a virtual environment"""
    # Check if we're already in a virtual environment
    if hasattr(sys, 'real_prefix') or (hasattr(sys, 'base_prefix') and sys.base_prefix != sys.prefix):
        print("✅ Using existing virtual environment")
        return True
    
    print("🔧 Setting up virtual environment...")
    
    # Determine platform-specific paths
    if platform.system() == 'Windows':
        python_executable = os.path.join(VENV_DIR, 'Scripts', 'python.exe')
        activate_script = os.path.join(VENV_DIR, 'Scripts', 'activate.bat')
    else:
        python_executable = os.path.join(VENV_DIR, 'bin', 'python')
        activate_script = os.path.join(VENV_DIR, 'bin', 'activate')
    
    # Create virtual environment if it doesn't exist
    if not os.path.exists(python_executable):
        try:
            subprocess.run([sys.executable, '-m', 'venv', VENV_DIR], check=True)
            print(f"✅ Virtual environment created at {VENV_DIR}")
        except Exception as e:
            print(f"❌ Failed to create virtual environment: {str(e)}")
            return False
    
    # Restart using the virtual environment's Python
    if os.path.exists(python_executable):
        print(f"🔄 Restarting with virtual environment Python: {python_executable}")
        os.execl(python_executable, python_executable, *sys.argv)
    
    return False

# Set up virtual environment before package installation
if not setup_virtual_environment():
    print("⚠️ Continuing without virtual environment - some features may be limited")

# --- Auto-Installer for Dependencies ---
def install_missing_packages():
    """Automatically install missing Python packages with fallbacks"""
    REQUIRED_PACKAGES = [
        "feedparser", "aiohttp", "beautifulsoup4", "google",
        "pyfiglet", "termcolor", "lxml", "pytz", "python-dateutil",
        "deep-translator", "nltk", "rich", "matplotlib", "wordcloud", "jinja2",
        "requests", "tqdm", "pygments", "art", "aiosqlite", 
        "cryptography", "pandas", "scikit-learn", "folium", "seaborn",
        "textblob", "geopy", "tweepy", "gtts", "reportlab",
        "python-docx", "dash", "plotly", "dash-bootstrap-components", "openpyxl"
    ]
    
    # Audio packages with platform-specific handling
    if platform.system() == 'Darwin':
        REQUIRED_PACKAGES.append("pyobjc")
    elif platform.system() == 'Windows':
        REQUIRED_PACKAGES.append("pywin32")
    
    OPTIONAL_PACKAGES = {
        "tweepy": "Social media monitoring will be limited",
        "geopy": "Geospatial features will be disabled",
        "folium": "Mapping features will be disabled",
        "gtts": "Audio briefings will be disabled",
        "reportlab": "PDF report generation will be disabled",
        "python-docx": "DOCX report generation will be disabled",
        "dash": "Live dashboard will be disabled",
        "openpyxl": "Excel export will be disabled"
    }
    
    # Try to import required packages, install if missing
    missing_packages = []
    for package in REQUIRED_PACKAGES:
        try:
            __import__(package.split('-')[0])
        except ImportError:
            missing_packages.append(package)
    
    if not missing_packages:
        print("✅ All required packages are installed")
        return True
    
    print(f"⚠️ Missing packages: {', '.join(missing_packages)}")
    
    # Determine if we're in a virtual environment
    in_venv = hasattr(sys, 'real_prefix') or (hasattr(sys, 'base_prefix') and sys.base_prefix != sys.prefix)
    
    # Prepare installation command
    if in_venv:
        cmd = [sys.executable, "-m", "pip", "install"]
    else:
        cmd = [sys.executable, "-m", "pip", "install", "--user"]
    
    # Try to install with standard pip
    try:
        print(f"⏳ Installing missing packages...")
        subprocess.check_call(cmd + missing_packages)
        print("✅ Packages installed successfully")
        return True
    except Exception as e:
        print(f"⚠️ Standard installation failed: {str(e)}")
    
    # Fallback to manual installation for critical packages
    print("⚠️ Attempting fallback installation...")
    success = True
    for package in missing_packages:
        try:
            # Special handling for problematic packages
            if package == "playsound":
                if platform.system() == 'Darwin':
                    subprocess.check_call(cmd + ["git+https://github.com/TaylorSMarks/playsound.git"])
                else:
                    subprocess.check_call(cmd + [package])
                
            print(f"  ✅ Installed {package}")
        except Exception as e:
            print(f"  ❌ Failed to install {package}: {str(e)}")
            if package in OPTIONAL_PACKAGES:
                print(f"  ⚠️ {OPTIONAL_PACKAGES[package]}")
            else:
                print(f"  ❗ Core feature may be affected")
                success = False
    
    # After installation, check if we can proceed
    if not success:
        print("⚠️ Some packages failed to install - the system may have limited functionality")
    
    return success

# Run installer before anything else
install_missing_packages()

# Now import all other modules with error handling
import feedparser
import aiohttp
from bs4 import BeautifulSoup
from cryptography.fernet import Fernet
from cryptography.hazmat.primitives import hashes
from cryptography.hazmat.primitives.kdf.pbkdf2 import PBKDF2HMAC
from cryptography.hazmat.backends import default_backend
import aiosqlite
from deep_translator import GoogleTranslator
import nltk
from nltk.tokenize import word_tokenize
from nltk.stem import PorterStemmer
from nltk.corpus import stopwords
from nltk.sentiment import SentimentIntensityAnalyzer
import matplotlib.pyplot as plt
import pandas as pd
import folium
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.cluster import KMeans
from termcolor import colored
import pyfiglet
from tqdm import tqdm
from gtts import gTTS
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
import dash
from dash import dcc, html
from dash.dependencies import Input, Output
import plotly.express as px
from reportlab.lib.styles import ParagraphStyle

# Platform-specific audio handling
try:
    if platform.system() == 'Darwin':
        from AppKit import NSSound
    elif platform.system() == 'Windows':
        import winsound
except ImportError:
    pass

# Download required NLTK data
try:
    nltk.download('punkt', quiet=True)
    nltk.download('stopwords', quiet=True)
    nltk.download('vader_lexicon', quiet=True)
except Exception as e:
    print(f"⚠️ NLP initialization failed: {str(e)}")
    stemmer = None
    stop_words = set()
    sia = None

# --- Enhanced Security Protocols ---
warnings.filterwarnings("ignore", category=UserWarning)
os.environ['PYTHONWARNINGS'] = 'ignore'

# --- Global Configuration ---
# Expanded list of global news sources
WEBSITE_LIST = [
    "https://www.yahoo.com/", "https://www.yahoo.co.jp/", "https://www.globo.com/",
    "https://www.nytimes.com/", "https://www.bbc.co.uk/", "https://www.detik.com/",
    "https://www.uol.com.br/", "https://www.bbc.com/", "https://www.indiatimes.com/",
    "https://www.hurriyet.com.tr/", "https://www.hindustantimes.com/", "https://www.ndtv.com/",
    "https://www.kumparan.com/", "https://www.msn.com/", "https://www.milliyet.com.tr/",
    "https://www.infobae.com/", "https://www.interia.pl/", "https://www.onet.pl/",
    "https://www.indianexpress.com/", "https://www.terra.com.br/", "https://www.kompas.com/",
    "https://www.conectate.com.do/", "https://www.thehindu.com/", "https://www.wp.pl/",
    "https://www.theguardian.com/", "https://www.liputan6.com/", "https://www.sanook.com/",
    "https://www.ouest-france.fr/", "https://www.liberation.fr/", "https://www.cnnbrasil.com.br/",
    "https://www.tribunnews.com/", "https://www.usatoday.com/", "https://www.cnn.com/",
    "https://www.tempo.co/", "https://www.foxnews.com/", "https://www.bild.de/",
    "https://www.lefigaro.fr/", "https://www.livehindustan.com/", "https://www.sabah.com.tr/",
    "https://www.dailymail.co.uk/", "https://www.thairath.co.th/", "https://www.elmundo.es/",
    "https://www.aajtak.in/", "https://www.news18.com/", "https://www.youm7.com/",
    "https://www.vnexpress.net/", "https://www.abplive.com/", "https://www.dzen.ru/",
    "https://www.usnews.com/", "https://www.ntv.com.tr/", "https://www.indiatoday.in/",
    "https://www.nhk.or.jp/", "https://www.haberturk.com/", "https://www.india.com/",
    "https://www.amarujala.com/", "https://www.24h.com.vn/", "https://www.corriere.it/",
    "https://www.kapook.com/", "https://www.lanacion.com.ar/", "https://www.business-standard.com/",
    "https://www.nikkei.com/", "https://www.cnnindonesia.com/", "https://www.baomoi.com/",
    "https://www.bhaskar.com/", "https://www.ukr.net/", "https://www.t-online.de/",
    "https://www.haberler.com/", "https://www.elpais.com/", "https://www.sondakika.com/",
    "https://www.jansatta.com/", "https://www.aol.com/", "https://www.repubblica.it/",
    "https://www.jagran.com/", "https://www.lavanguardia.com/", "https://www.ndr.de/",
    "https://www.cbsnews.com/", "https://www.thesun.co.uk/", "https://www.bfmtv.com/",
    "https://www.spiegel.de/", "https://www.buzzfeed.com/", "https://www.nbcnews.com/",
    "https://www.nypost.com/", "https://www.ndtv.in/", "https://www.aljazeera.com/",
    "https://www.tagesschau.de/", "https://www.cnnturk.com/", "https://www.estadao.com.br/",
    "https://www.abc.es/", "https://www.n-tv.de/", "https://www.rtve.es/",
    "https://www.rambler.ru/", "https://www.apnews.com/", "https://www.lemonde.fr/",
    "https://www.abc.net.au/", "https://www.almasryalyoum.com/", "https://www.iltalehti.fi/",
    "https://www.clarin.com/", "https://www.reuters.com/", "https://www.haber7.com/",
    "https://www.antaranews.com/"
]

# --- Logging Setup ---
def setup_logging():
    """Configure secure logging system"""
    log_file = os.path.join(LOG_DIR, "mgcis_operations.log")
    
    logger = logging.getLogger("MG_CIS")
    logger.setLevel(logging.INFO)
    
    # File handler with rotation
    file_handler = RotatingFileHandler(
        log_file, maxBytes=5*1024*1024, backupCount=5
    )
    file_handler.setLevel(logging.INFO)
    
    # Console handler
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.WARNING)
    
    # Formatting
    formatter = logging.Formatter(
        '%(asctime)s - %(name)s - %(levelname)s - %(message)s',
        datefmt='%Y-%m-%d %H:%M:%S'
    )
    file_handler.setFormatter(formatter)
    console_handler.setFormatter(formatter)
    
    logger.addHandler(file_handler)
    logger.addHandler(console_handler)
    
    return logger

logger = setup_logging()

# --- Configuration Management ---
CONFIG_FILE = os.path.join(CONFIG_DIR, "mgcis_config.ini")

def load_config():
    """Load or create configuration file with robust defaults"""
    # Define default configuration with version
    DEFAULT_CONFIG = {
        'encryption_key': base64.urlsafe_b64encode(os.urandom(32)).decode(),
        'threat_scan_interval': '3600',
        'alert_threshold': '80',
        'max_concurrent_requests': '15',
        'social_media_enabled': 'True',
        'geospatial_enabled': 'True',
        'llm_integration': 'False',
        'llm_api_key': '',
        'twitter_api_key': '',
        'twitter_api_secret': '',
        'twitter_access_token': '',
        'twitter_access_secret': '',
        'twitter_bearer_token': '',
        'audio_enabled': 'True',
        'audio_voice': 'f1',  # Female voice 1
        'audio_language': 'en',
        'version': '2.0.0'  # Updated version
    }
    
    config = configparser.ConfigParser()
    
    if os.path.exists(CONFIG_FILE):
        config.read(CONFIG_FILE)
        
        # Fix missing keys by adding default values
        for key, value in DEFAULT_CONFIG.items():
            if not config.has_option('DEFAULT', key):
                config.set('DEFAULT', key, value)
                print(f"🔧 Added missing config key: {key}")
                
        # Save updated config
        with open(CONFIG_FILE, 'w') as configfile:
            config.write(configfile)
    else:
        # Create default configuration
        config['DEFAULT'] = DEFAULT_CONFIG
        with open(CONFIG_FILE, 'w') as configfile:
            config.write(configfile)
        os.chmod(CONFIG_FILE, 0o600)
    
    return config

config = load_config()

# --- Enhanced Crypto Setup ---
# Use persistent key from config instead of generating from system info
if config['DEFAULT'].get('encryption_key'):
    CRYPTO_KEY = config['DEFAULT']['encryption_key'].encode()
else:
    # Generate new key and save to config
    new_key = base64.urlsafe_b64encode(os.urandom(32)).decode()
    config['DEFAULT']['encryption_key'] = new_key
    with open(CONFIG_FILE, 'w') as configfile:
        config.write(configfile)
    CRYPTO_KEY = new_key.encode()

cipher_suite = Fernet(CRYPTO_KEY)

# ======================
# GLOBAL CONFIGURATION
# ======================
# Database setup
DB_FILE = os.path.join(DB_DIR, "war_intel.db")

# Predefined RSS feeds (encrypted storage)
FEED_STORAGE = os.path.join(DB_DIR, "feeds.enc")

# Enhanced keyword groups with weights
KEYWORD_GROUPS = {
    "critical": {
        "war": 20, "attack": 18, "nuclear": 25, "terrorism": 18, "invasion": 20, 
        "strike": 16, "casualties": 15, "drone": 14, "russia": 12, "ukraine": 12, 
        "israel": 12, "iran": 12, "conflict": 15, "bomb": 18, "missile": 16, 
        "offensive": 15, "combat": 14, "death": 15, "died": 14, "fatality": 15,
        "fatalities": 15, "killed": 16, "victim": 13, "bombing": 18, "explosion": 17,
        "explosive": 16, "detonation": 16, "airstrike": 17, "counterstrike": 16,
        "uav": 14, "artillery": 14, "bombardment": 16, "counteroffensive": 16,
        "militia": 13, "nato": 12, "atomic": 22, "thermonuclear": 25,
        "chemical": 24, "biological": 24, "radiological": 23
    },
    "military": {
        "army": 8, "navy": 8, "air force": 9, "special forces": 10, "deployment": 12, 
        "battalion": 10, "regiment": 10, "troops": 11, "military doctrine": 9, 
        "rules of engagement": 10, "war cabinet": 12, "armed forces": 9, "command": 8, 
        "control": 8, "base": 9, "fortification": 9, "outpost": 9, "mobilization": 12,
        "guerrilla": 10, "counterinsurgency": 11, "occupation": 13, "siege": 14, 
        "blockade": 13, "raid": 12, "uprising": 12, "rebellion": 11, "insurgency": 12,
        "revolution": 13, "coup": 15, "martial law": 14, "state of emergency": 13,
        "amphibious": 11, "armored": 10, "artillery": 11, "logistics": 9
    },
    "cyber": {
        "cyberwar": 15, "hack": 12, "sigint": 10, "digital battlefield": 14, 
        "data breach": 13, "cyberwarfare": 15, "information operations": 12, 
        "electronic warfare": 13, "psychological operations": 11, "psyop": 10, 
        "uav operations": 11, "precision strike": 12, "satellite surveillance": 11,
        "space domain awareness": 10, "space warfare": 13, "cyber-physical systems": 9,
        "cyber-physical warfare": 12, "information warfare": 14, "hybrid warfare": 13,
        "cyber espionage": 14, "signal intelligence": 11, "electronic surveillance": 12,
        "cyber operations": 13, "zero-day": 16, "apt": 15, "ransomware": 14
    },
    "intel": {
        "national security": 12, "defense policy": 10, "security council": 9, "nsa": 10, 
        "cia": 10, "dod": 9, "fbi": 9, "mossad": 10, "kgb": 10, "mi6": 10, "fsb": 10, 
        "covert operation": 14, "espionage": 13, "intelligence": 11, "humint": 9, "sigint": 9, 
        "clandestine": 12, "mole": 13, "double agent": 12, "defector": 11, "sleeper cell": 13, 
        "surveillance": 11, "counterintelligence": 12, "covert agent": 12, "operative": 11,
        "informant": 10, "counterterrorism": 13, "reconnaissance": 11, "infiltration": 12,
        "secret service": 10, "black ops": 14, "covert action": 13, "intelligence briefing": 9,
        "tradecraft": 8, "classified": 10, "declassified": 9, "asset": 11, "exfiltration": 12,
        "dead drop": 11, "brush pass": 10
    },
    "weapons": {
        "f-35": 12, "hypersonic": 15, "ballistic": 14, "cruise": 13, "s-400": 11,
        "patriot": 10, "ak-47": 9, "tank": 10, "submarine": 11, "aircraft carrier": 12,
        "icbm": 16, "tactical nuke": 18, "drone swarm": 15, "laser weapon": 14,
        "railgun": 13, "cyber weapon": 14, "biological weapon": 19, "chemical weapon": 19
    }
}

# Hashtags and keywords for social media monitoring
HASHTAGS = [
    "#BreakingNews", "#WarNews", "#ConflictZone", "#MilitaryUpdate", "#GlobalCrisis",
    "#HumanitarianCrisis", "#WarCrimes", "#Ceasefire", "#PeaceTalks", "#UkraineWar",
    "#RussiaUkraineConflict", "#StandWithUkraine", "#KyivUnderAttack", "#PutinWar",
    "#UkraineInvasion", "#SlavaUkraini", "#IsraelIranWar", "#OperationRisingLion",
    "#MiddleEastCrisis", "#IsraelStrikesIran", "#IranRetaliation", "#LiveUpdate",
    "#IndiaVsPakistan", "#OperationSindoor", "#IndianArmy", "#DefendIndia",
    "#PakistanConflict", "#PehelgamAttack", "#CyberWar", "#DroneStrike",
    "#HybridWarfare", "#InformationWarfare", "#DigitalBattlefield"
]

KEYWORDS = [
    "National Security", "Defense Policy", "Military Doctrine", "Rules of Engagement",
    "State of Emergency", "Martial Law", "War Cabinet", "Security Council",
    "Preemptive Strike", "Geopolitical Tensions", "Strategic Competition",
    "Political Warfare", "International Armed Conflict", "Non-International Armed Conflict",
    "Civil War", "Insurgency", "Occupation", "Proxy War", "Hybrid Warfare",
    "Grey-Zone Conflict", "Guerrilla Warfare", "Counterinsurgency",
    "Low-Intensity Conflict", "Geneva Conventions", "International Humanitarian Law",
    "War Crimes", "Crimes Against Humanity", "Combatant Status", "Belligerent Party",
    "Humanitarian Corridor", "Ceasefire Agreement", "Peacekeeping Mission",
    "Blue Helmet", "UN Peacekeeping", "Humanitarian Aid", "Relief Efforts",
    "Refugee Crisis", "Displaced Persons", "Civilian Casualties", "War Torn",
    "Reconstruction Efforts", "Post-Conflict Reconstruction", "Cyberwarfare",
    "Information Operations", "Drone Strike", "Electronic Warfare",
    "Psychological Operations", "PSYOP", "Digital Battlefield",
    "Unmanned Aerial Vehicle", "UAV Operations", "Precision Strike",
    "Satellite Surveillance", "Space Domain Awareness", "Space Warfare",
    "Cyber-Physical Systems", "Cyber-Physical Warfare", "war", "civil war",
    "world war", "cold war", "holy war", "limited war", "conflict", "battle",
    "combat", "fighting", "warfare", "hostilities", "engagement", "offensive",
    "counterattack", "skirmish", "clash", "incursion", "siege", "blockade",
    "ambush", "assault", "police action", "conflagration", "death", "died",
    "fatality", "fatalities", "killed", "victim", "casualties", "bomb", "bombing",
    "explosion", "explosive", "detonation", "strike", "airstrike", "missile",
    "counterstrike", "drone", "UAV", "nuclear", "atomic", "thermonuclear",
    "artillery", "bombardment", "regiment", "division", "battalion", "brigade",
    "encircled", "encirclement", "front lines", "battle lines", "retreat",
    "withdrawal", "deployment", "mobilization", "counteroffensive", "raid",
    "uprising", "rebellion", "insurgency", "revolution", "coup", "riot",
    "military action", "armed forces", "ARMY", "NAVY", "AIR FORCE", "marines",
    "special forces", "command", "control", "base", "fortification", "outpost",
    "escalation", "de-escalation", "proxy war", "asymmetric warfare",
    "hybrid warfare", "psychological operations", "propaganda", "targeted strike",
    "casualty count", "terrorism", "attack", "militia", "NSA", "CIA", "DOD", "FBI",
    "Mossad", "KGB", "MI6", "FSB", "covert operation", "espionage", "intelligence",
    "HUMINT", "SIGINT", "cyber espionage", "clandestine", "mole", "double agent",
    "defector", "sleeper cell", "surveillance", "counterintelligence",
    "covert agent", "operative", "informant", "counterterrorism", "reconnaissance",
    "signal intelligence", "electronic surveillance", "infiltration",
    "cyber operations", "data breach", "secret service", "black ops",
    "covert action", "intelligence briefing", "tradecraft", "classified",
    "declassified", "asset", "exfiltration", "cyber warfare", "NASA", "UFO"
]

# Global threat levels
THREAT_LEVELS = {
    0: ("🟢", "Low", "No immediate threats detected"),
    30: ("🟡", "Guarded", "Potential threats developing"),
    50: ("🟠", "Elevated", "Significant threat activity"),
    70: ("🔴", "High", "Critical threats imminent"),
    90: ("🛑", "Severe", "Active conflict situations")
}

HISTORICAL_CONFLICTS = {
    "Israel-Palestine": {"start": "1948", "casualties": "100,000+", "status": "Ongoing"},
    "India-Pakistan": {"start": "1947", "casualties": "100,000+", "status": "Intermittent"},
    "Russia-Ukraine": {"start": "2014", "casualties": "500,000+", "status": "Active"},
    "US-Iran": {"start": "1979", "casualties": "500,000+", "status": "Tensions"},
    "China-Taiwan": {"start": "1949", "casualties": "1,000,000+", "status": "Tensions"},
    "Korea": {"start": "1950", "casualties": "2,500,000+", "status": "Ceasefire"},
    "Syria": {"start": "2011", "casualties": "600,000+", "status": "Ongoing"},
    "Afghanistan": {"start": "1978", "casualties": "2,000,000+", "status": "Active"},
    "Iraq": {"start": "2003", "casualties": "1,000,000+", "status": "Ongoing"},
    "Yemen": {"start": "2014", "casualties": "377,000+", "status": "Ongoing"},
    "Ethiopia": {"start": "2020", "casualties": "500,000+", "status": "Ongoing"},
    "India-China": {"start": "1962", "casualties": "10,000+", "status": "Border tensions"},
    "China-USA": {"start": "1949", "casualties": "100,000+", "status": "Trade war"},
    "Iraq-USA": {"start": "2003", "casualties": "500,000+", "status": "Post-conflict"},
    "Israel-Iraq": {"start": "1981", "casualties": "1,000+", "status": "Historical"},
    "Russia-Georgia": {"start": "2008", "casualties": "1,000+", "status": "Resolved"},
    "Vietnam War": {"start": "1955", "casualties": "3,000,000+", "status": "Historical"},
    "Gulf War": {"start": "1990", "casualties": "150,000+", "status": "Historical"},
    "Falklands War": {"start": "1982", "casualties": "1,000+", "status": "Historical"},
    "Korean War": {"start": "1950", "casualties": "2,500,000+", "status": "Ceasefire",
    }
}

CONFLICT_GROUPS = {
    "Israel-Palestine": re.compile(r"Israel|Palestine|Gaza|West Bank|Hamas|IDF", re.IGNORECASE),
    "India-Pakistan": re.compile(r"India|Pakistan|Kashmir|LoC", re.IGNORECASE),
    "Russia-Ukraine": re.compile(r"Russia|Ukraine|Kyiv|Moscow|Donbas|Zelensky|Putin", re.IGNORECASE),
    "US-Iran": re.compile(r"USA|US|Iran|Tehran|Persian Gulf|Strait of Hormuz", re.IGNORECASE),
    "China-Taiwan": re.compile(r"China|Taiwan|Beijing|Taipei|TSMC|South China Sea", re.IGNORECASE),
    "India-China": re.compile(r"India|China|Ladakh|Galwan|Doklam|Line of Actual Control", re.IGNORECASE),
    "China-USA": re.compile(r"China|USA|US|trade war|technology war|South China Sea", re.IGNORECASE),
    "Iraq-USA": re.compile(r"Iraq|USA|US|Baghdad|Saddam Hussein|Operation Iraqi Freedom", re.IGNORECASE),
    "Israel-Iraq": re.compile(r"Israel|Iraq|Osirak|Operation Opera", re.IGNORECASE),
    "Global Conflicts": re.compile(r"war|conflict|battle|attack|strike", re.IGNORECASE),
    "Middle East": re.compile(r"Syria|Yemen|Lebanon|Iraq|Afghanistan|Turkey", re.IGNORECASE),
    "Africa": re.compile(r"Ethiopia|Sudan|Somalia|Mali|Niger|Burkina Faso", re.IGNORECASE),
    "Asia-Pacific": re.compile(r"Korea|Japan|Australia|Philippines|South China Sea", re.IGNORECASE),
    "Europe": re.compile(r"Germany|France|UK|Poland|Baltic|NATO", re.IGNORECASE)
}

# ======================
# TEXT PROCESSING SETUP
# ======================
try:
    stemmer = PorterStemmer()
    stop_words = set(stopwords.words('english'))
    sia = SentimentIntensityAnalyzer()
except Exception as e:
    print(f"⚠️ NLP initialization failed: {str(e)}")
    stemmer = None
    stop_words = set()
    sia = None

# ======================
# DATABASE FUNCTIONS
# ======================
async def init_db():
    """Initialize database with proper directory creation"""
    db_path = os.path.join(DB_DIR, "war_intel.db")
    try:
        async with aiosqlite.connect(db_path) as db:
            await db.execute("""
            CREATE TABLE IF NOT EXISTS articles (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT,
                url TEXT UNIQUE,
                source TEXT,
                published_date TEXT,
                summary TEXT,
                casualties_killed INTEGER DEFAULT 0,
                casualties_injured INTEGER DEFAULT 0,
                casualties_civilian INTEGER DEFAULT 0,
                conflict_group TEXT,
                author TEXT,
                entities TEXT,
                sentiment REAL,
                threat_score REAL,
                matched_terms TEXT,
                raw TEXT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP,
                language TEXT DEFAULT 'en'
            )
            """)
            
            await db.execute("""
            CREATE TABLE IF NOT EXISTS daily_stats (
                date TEXT PRIMARY KEY,
                killed INTEGER DEFAULT 0,
                injured INTEGER DEFAULT 0,
                strikes INTEGER DEFAULT 0,
                civilian_casualties INTEGER DEFAULT 0,
                conflict_count INTEGER DEFAULT 0
            )
            """)
            
            await db.execute("""
            CREATE TABLE IF NOT EXISTS alerts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT,
                url TEXT,
                conflict_group TEXT,
                severity TEXT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
            )
            """)
            
            await db.execute("""
            CREATE TABLE IF NOT EXISTS conflicts (
                name TEXT PRIMARY KEY,
                start_date TEXT,
                casualties TEXT,
                status TEXT,
                last_updated TEXT
            )
            """)
            
            await db.execute("""
            CREATE TABLE IF NOT EXISTS tools (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                name TEXT UNIQUE,
                description TEXT,
                code TEXT,
                created TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
            """)
            
            await db.execute("""
            CREATE TABLE IF NOT EXISTS social_media (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                platform TEXT,
                author TEXT,
                content TEXT,
                url TEXT,
                post_timestamp DATETIME,
                keywords TEXT,
                sentiment REAL,
                conflict_group TEXT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
            )
            """)
            
            # Add high-priority alerts table
            await db.execute("""
            CREATE TABLE IF NOT EXISTS high_priority_alerts (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                title TEXT,
                url TEXT,
                source TEXT,
                threat_score REAL,
                conflict_group TEXT,
                matched_terms TEXT,
                timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
            )
            """)
            
            # Insert historical conflicts
            for name, data in HISTORICAL_CONFLICTS.items():
                await db.execute("""
                INSERT OR IGNORE INTO conflicts (name, start_date, casualties, status)
                VALUES (?, ?, ?, ?)
                """, (name, data["start"], data["casualties"], data["status"]))
            
            await db.commit()
        return True
    except Exception as e:
        logger.error(f"Database initialization failed: {str(e)}")
        return False

async def get_db():
    """Get database connection with connection pooling"""
    db_path = os.path.join(DB_DIR, "war_intel.db")
    try:
        return await aiosqlite.connect(db_path)
    except Exception as e:
        logger.error(f"Database connection failed: {str(e)}")
        return None

# ======================
# AUDIO BRIEFING SYSTEM
# ======================
class AudioBriefing:
    def __init__(self):
        self.playback_queue = queue.Queue()
        self.currently_playing = False
        self.playback_thread = threading.Thread(target=self._playback_worker, daemon=True)
        self.playback_thread.start()
        self.playback_lock = threading.Lock()
        
        # Enhanced voice options with accents
        self.voice_options = {
            'f1': ('en', 'female', 'Sophia (US English)'),
            'f2': ('en-au', 'female', 'Ava (Australian)'),
            'f3': ('en', 'female', 'English (US)'),  # Fixed from unsupported 'en-in'
            'f4': ('en-gb', 'female', 'Emma (British)'),
            'f5': ('fr', 'female', 'Camille (French)'),
            'f6': ('de', 'female', 'Vicki (German)'),
            'f7': ('it', 'female', 'Carla (Italian)'),
            'f8': ('es', 'female', 'Conchita (Spanish)'),
            'm1': ('en', 'male', 'James (US English)'),
            'm2': ('en-gb', 'male', 'Harry (British)')
        }
        
        self.enabled = config['DEFAULT'].getboolean('audio_enabled', True)
        self.voice = config['DEFAULT'].get('audio_voice', 'f1')
        self.language = config['DEFAULT'].get('audio_language', 'en')
        
    def _playback_worker(self):
        """Background worker to play audio files sequentially"""
        while True:
            audio_path = self.playback_queue.get()
            if audio_path is None:  # Termination signal
                break
                
            try:
                self._play_audio(audio_path)
            except Exception as e:
                logger.error(f"Audio playback error: {str(e)}")
            finally:
                self.playback_queue.task_done()
    
    def _play_audio(self, audio_path):
        """Play audio briefing using platform-specific methods"""
        if not audio_path or not os.path.exists(audio_path):
            return False
            
        try:
            if platform.system() == 'Darwin':
                # macOS
                from AppKit import NSSound
                sound = NSSound.alloc()
                sound.initWithContentsOfFile_byReference_(audio_path, True)
                sound.play()
                # Wait for playback to finish
                while sound.isPlaying():
                    time.sleep(0.1)
                return True
            elif platform.system() == 'Windows':
                # Windows
                import winsound
                winsound.PlaySound(audio_path, winsound.SND_FILENAME | winsound.SND_ASYNC)
                # Wait for playback to finish
                time.sleep(self._get_audio_duration(audio_path))
                return True
            else:
                # Linux/other - complete implementation
                if shutil.which("paplay"):  # PulseAudio
                    subprocess.run(["paplay", audio_path])
                    return True
                elif shutil.which("aplay"):  # ALSA
                    subprocess.run(["aplay", audio_path])
                    return True
                elif shutil.which("mpg123"):  # MP3 player
                    subprocess.run(["mpg123", audio_path])
                    return True
                else:
                    logger.warning("No suitable audio player found on Linux")
                    return False
        except Exception as e:
            logger.error(f"Audio playback failed: {str(e)}")
            return False
            
    def _get_audio_duration(self, audio_path):
        """Estimate audio duration (simplified)"""
        # 1 second per 16KB as approximation
        return max(1, os.path.getsize(audio_path) // 16000)
    
    def generate_briefing(self, text, filename="briefing.mp3", brief_type="standard"):
        """Generate audio briefing from text with enhanced error handling"""
        if not self.enabled or not text:
            return None
            
        try:
            voice_settings = self.voice_options.get(self.voice, self.voice_options['f1'])
            lang, gender, _ = voice_settings
            
            # Handle special cases for certain languages/accents
            lang = 'en' if lang == 'en-in' else lang  # Fallback for unsupported Indian English
            
            tts = gTTS(
                text=text,
                lang=lang,
                slow=False,
                lang_check=False
            )
            
            audio_path = os.path.join(AUDIO_DIR, filename)
            tts.save(audio_path)
            return audio_path
        except Exception as e:
            logger.error(f"Audio briefing generation failed: {str(e)}")
            # Fallback to default voice
            if self.voice != 'f1':
                logger.info("Attempting fallback to default voice")
                try:
                    tts = gTTS(
                        text=text,
                        lang='en',
                        slow=False,
                        lang_check=False
                    )
                    audio_path = os.path.join(AUDIO_DIR, filename)
                    tts.save(audio_path)
                    return audio_path
                except Exception as fallback_e:
                    logger.error(f"Fallback audio generation failed: {str(fallback_e)}")
            return None
            
    def play_briefing(self, audio_path):
        """Queue audio briefing for playback"""
        if not audio_path or not os.path.exists(audio_path):
            return False
            
        with self.playback_lock:
            self.playback_queue.put(audio_path)
        return True
        
    def generate_critical_briefing(self, alerts):
        """Generate briefing for critical alerts"""
        if not alerts or len(alerts) == 0:
            return None
            
        # Prepare briefing text
        briefing = "Global Threat Briefing. "
        max_score = max(a['score'] for a in alerts) if alerts else 0
        threat_level = self.get_threat_level_name(max_score)
        briefing += f"Current threat level: {threat_level}. "
        briefing += f"Detected {len(alerts)} critical alerts. "
        
        # Group by conflict
        conflict_groups = defaultdict(list)
        for alert in alerts:
            conflict_groups[alert.get('conflict_group', 'Global')].append(alert)
        
        # Summarize each conflict
        for conflict, group in conflict_groups.items():
            highest_alert = max(group, key=lambda x: x['score'])
            briefing += f"In {conflict}, highest threat: {highest_alert['title']}. "
            
        # Add casualty information
        casualties = sum(a.get('casualties_killed', 0) for a in alerts)
        if casualties > 0:
            briefing += f"Total casualties reported: {casualties}. "
            
        briefing += "End of briefing."
        return briefing
        
    def generate_dashboard_briefing(self, intel_data):
        """Generate audio briefing for dashboard"""
        if not intel_data:
            return "No intelligence data available."
            
        max_score = max(a['score'] for a in intel_data) if intel_data else 0
        threat_level = self.get_threat_level_name(max_score)
        
        summary = f"Live Intelligence Dashboard Update. Current global threat level is {threat_level}. "
        summary += f"Detected {len(intel_data)} active alerts. "
        
        # Get casualty totals
        total_killed = sum(a.get('casualties_killed', 0) for a in intel_data)
        total_injured = sum(a.get('casualties_injured', 0) for a in intel_data)
        
        if total_killed > 0:
            summary += f"Total casualties reported: {total_killed} killed, {total_injured} injured. "
        
        # Get top conflicts
        conflict_counts = defaultdict(int)
        for alert in intel_data:
            conflict_counts[alert['conflict_group']] += 1
            
        top_conflicts = sorted(conflict_counts.items(), key=lambda x: x[1], reverse=True)[:3]
        if top_conflicts:
            summary += "Primary conflict zones: "
            summary += ", ".join([f"{conflict[0]} with {conflict[1]} alerts" for conflict in top_conflicts]) + ". "
        
        # Get highest alert
        highest_alert = max(intel_data, key=lambda x: x['score'], default=None)
        if highest_alert:
            summary += f"Highest priority alert: {highest_alert['title']} in {highest_alert['conflict_group']}. "
            
        summary += "End of briefing."
        return summary
        
    def get_threat_level_name(self, score):
        """Get threat level name based on score"""
        for threshold, (_, name, _) in sorted(THREAT_LEVELS.items(), reverse=True):
            if score >= threshold:
                return name
        return "Unknown"
        
    def select_voice(self):
        """Select voice for audio briefings"""
        print("\n🎤 SELECT VOICE FOR AUDIO BRIEFINGS")
        print("=" * 50)
        for i, (voice_id, (_, _, name)) in enumerate(self.voice_options.items(), 1):
            print(f"{i}. {name}")
            
        choice = input("\nSelect voice (1-10): ")
        try:
            idx = int(choice) - 1
            if 0 <= idx < len(self.voice_options):
                voice_id = list(self.voice_options.keys())[idx]
                config['DEFAULT']['audio_voice'] = voice_id
                with open(CONFIG_FILE, 'w') as configfile:
                    config.write(configfile)
                print(f"✅ Voice set to: {self.voice_options[voice_id][2]}")
            else:
                print("❌ Invalid selection")
        except:
            print("❌ Invalid input")
            
    def play_critical_update(self, alerts):
        """Play a critical update briefing"""
        if not alerts:
            return False
            
        # Get the top alert
        top_alert = max(alerts, key=lambda x: x['score'])
        threat_level = self.get_threat_level_name(top_alert['score'])
        briefing_text = f"CRITICAL UPDATE: {top_alert['title']} in {top_alert['conflict_group']}. "
        briefing_text += f"Threat level: {threat_level}. "
        if top_alert.get('casualties_killed', 0) > 0:
            briefing_text += f"Casualties: {top_alert['casualties_killed']} killed. "
            
        audio_file = self.generate_briefing(briefing_text, "critical_update.mp3")
        if audio_file:
            return self.play_briefing(audio_file)
        return False
        
    def play_complete_briefing(self, alerts):
        """Play a complete briefing of all critical alerts"""
        if not alerts:
            return False
            
        briefing_text = self.generate_critical_briefing(alerts)
        audio_file = self.generate_briefing(briefing_text, "complete_briefing.mp3")
        if audio_file:
            return self.play_briefing(audio_file)
        return False

# ======================
# REPORT GENERATION
# ======================
class ReportGenerator:
    def __init__(self):
        self.audio = AudioBriefing()
        
    def generate_detailed_report(self, report_data, format="json"):
        """Generate detailed intelligence report in multiple formats"""
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        base_filename = f"intel_report_{timestamp}"
        
        # Generate JSON report as base
        json_report = self._generate_json_report(report_data, base_filename)
        
        # Generate in requested format
        if format == "pdf":
            return self._generate_pdf_report(report_data, base_filename)
        elif format == "docx":
            return self._generate_docx_report(report_data, base_filename)
        elif format == "audio":
            return self._generate_audio_briefing(report_data, base_filename)
        elif format == "csv":
            return self._generate_csv_report(report_data, base_filename)
        elif format == "excel":
            return self._generate_excel_report(report_data, base_filename)
        else:
            return json_report
            
    def _generate_json_report(self, report_data, base_filename):
        """Generate JSON report"""
        report_file = os.path.join(REPORT_DIR, f"{base_filename}.json")
        with open(report_file, 'w') as f:
            json.dump(report_data, f, indent=2, cls=DateTimeEncoder)
        return report_file
        
    def _generate_pdf_report(self, report_data, base_filename):
        """Generate PDF report with detailed formatting"""
        try:
            report_file = os.path.join(REPORT_DIR, f"{base_filename}.pdf")
            doc = SimpleDocTemplate(report_file, pagesize=letter)
            styles = getSampleStyleSheet()
            elements = []
            
            # Title
            title_style = ParagraphStyle(
                name='Title',
                fontSize=18,
                alignment=1,
                spaceAfter=12
            )
            elements.append(Paragraph("MILITARY-GLOBAL CONFLICT INTELLIGENCE REPORT", title_style))
            elements.append(Paragraph(f"Generated: {report_data['timestamp']}", styles['Normal']))
            elements.append(Spacer(1, 12))
            
            # Threat level
            threat_level = report_data['threat_level']
            threat_text = f"Global Threat Level: {threat_level[1]} - {threat_level[2]}"
            elements.append(Paragraph(threat_text, styles['Heading2']))
            elements.append(Spacer(1, 12))
            
            # Executive summary
            exec_summary = self._generate_executive_summary(report_data)
            elements.append(Paragraph("Executive Summary", styles['Heading2']))
            elements.append(Paragraph(exec_summary, styles['Normal']))
            elements.append(Spacer(1, 12))
            
            # Daily statistics
            stats = report_data['daily_stats']
            stat_data = [
                ['Metric', 'Value'],
                ['Killed', stats['killed']],
                ['Injured', stats['injured']],
                ['Civilian Casualties', stats['civilian_casualties']],
                ['Strikes/Attacks', stats['strikes']],
                ['Active Conflicts', stats['conflict_count']]
            ]
            
            stat_table = Table(stat_data)
            stat_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 1, colors.black)
            ]))
            
            elements.append(Paragraph("Daily Statistics", styles['Heading2']))
            elements.append(stat_table)
            elements.append(Spacer(1, 12))
            
            # Conflict details
            elements.append(Paragraph("Conflict Zone Analysis", styles['Heading2']))
            for conflict, killed, injured, reports in report_data['conflict_stats']:
                conflict_text = f"{conflict}: {killed} killed, {injured} injured across {reports} incidents"
                elements.append(Paragraph(conflict_text, styles['Normal']))
            elements.append(Spacer(1, 12))
            
            # Top alerts
            elements.append(Paragraph("Top Threat Alerts", styles['Heading2']))
            for i, alert in enumerate(report_data['top_alerts'], 1):
                alert_text = f"{i}. {alert['title']} (Score: {alert['score']})"
                elements.append(Paragraph(alert_text, styles['Normal']))
                elements.append(Paragraph(f"Source: {alert['source']} | {alert['published']}", styles['Italic']))
                elements.append(Spacer(1, 6))
            
            # Build PDF
            doc.build(elements)
            return report_file
        except Exception as e:
            logger.error(f"PDF report generation failed: {str(e)}")
            return self._generate_json_report(report_data, base_filename)
            
    def _generate_docx_report(self, report_data, base_filename):
        """Generate DOCX report"""
        try:
            report_file = os.path.join(REPORT_DIR, f"{base_filename}.docx")
            doc = Document()
            
            # Title
            doc.add_heading('MILITARY-GLOBAL CONFLICT INTELLIGENCE REPORT', 0)
            doc.add_paragraph(f"Generated: {report_data['timestamp']}")
            doc.add_paragraph()
            
            # Threat level
            threat_level = report_data['threat_level']
            doc.add_heading(f"Global Threat Level: {threat_level[1]}", level=1)
            doc.add_paragraph(threat_level[2])
            doc.add_paragraph()
            
            # Executive summary
            exec_summary = self._generate_executive_summary(report_data)
            doc.add_heading('Executive Summary', level=1)
            doc.add_paragraph(exec_summary)
            doc.add_paragraph()
            
            # Daily statistics
            doc.add_heading('Daily Statistics', level=1)
            stats = report_data['daily_stats']
            table = doc.add_table(rows=6, cols=2)
            table.style = 'LightShading'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = 'Value'
            
            # Data rows
            metrics = [
                ('Killed', stats['killed']),
                ('Injured', stats['injured']),
                ('Civilian Casualties', stats['civilian_casualties']),
                ('Strikes/Attacks', stats['strikes']),
                ('Active Conflicts', stats['conflict_count'])
            ]
            
            for i, (metric, value) in enumerate(metrics, 1):
                row_cells = table.rows[i].cells
                row_cells[0].text = metric
                row_cells[1].text = str(value)
            
            doc.add_paragraph()
            
            # Conflict details
            doc.add_heading('Conflict Zone Analysis', level=1)
            for conflict, killed, injured, reports in report_data['conflict_stats']:
                doc.add_paragraph(
                    f"{conflict}: {killed} killed, {injured} injured across {reports} incidents",
                    style='ListBullet'
                )
            doc.add_paragraph()
            
            # Top alerts
            doc.add_heading('Top Threat Alerts', level=1)
            for i, alert in enumerate(report_data['top_alerts'], 1):
                doc.add_paragraph(f"{i}. {alert['title']}", style='ListNumber')
                doc.add_paragraph(f"Threat Score: {alert['score']}", style='BodyText')
                doc.add_paragraph(f"Source: {alert['source']} | Published: {alert['published']}", style='BodyText')
                doc.add_paragraph()
            
            doc.save(report_file)
            return report_file
        except Exception as e:
            logger.error(f"DOCX report generation failed: {str(e)}")
            return self._generate_json_report(report_data, base_filename)
            
    def _generate_audio_briefing(self, report_data, base_filename):
        """Generate audio briefing"""
        briefing_text = self._generate_executive_summary(report_data)
        audio_file = self.audio.generate_briefing(briefing_text, f"{base_filename}.mp3")
        if audio_file:
            self.audio.play_briefing(audio_file)
        return audio_file
    
    def _generate_csv_report(self, report_data, base_filename):
        """Generate CSV report"""
        try:
            report_file = os.path.join(REPORT_DIR, f"{base_filename}.csv")
            
            # Prepare CSV content
            csv_lines = [
                "MILITARY-GLOBAL CONFLICT INTELLIGENCE REPORT",
                f"Generated: {report_data['timestamp']}",
                ""
            ]
            
            # Threat level
            threat_level = report_data['threat_level']
            csv_lines.append(f"Global Threat Level:,{threat_level[1]}")
            csv_lines.append(f"Description:,{threat_level[2]}")
            csv_lines.append("")
            
            # Daily stats
            stats = report_data['daily_stats']
            csv_lines.append("Daily Statistics")
            csv_lines.append("Metric,Value")
            csv_lines.append(f"Killed,{stats['killed']}")
            csv_lines.append(f"Injured,{stats['injured']}")
            csv_lines.append(f"Civilian Casualties,{stats['civilian_casualties']}")
            csv_lines.append(f"Strikes/Attacks,{stats['strikes']}")
            csv_lines.append(f"Active Conflicts,{stats['conflict_count']}")
            csv_lines.append("")
            
            # Conflict stats
            csv_lines.append("Conflict Zone Analysis")
            csv_lines.append("Conflict Zone,Killed,Injured,Incidents")
            for conflict, killed, injured, reports in report_data['conflict_stats']:
                csv_lines.append(f"{conflict},{killed},{injured},{reports}")
            csv_lines.append("")
            
            # Top alerts
            csv_lines.append("Top Threat Alerts")
            csv_lines.append("Rank,Title,Threat Score,Source,Published,Conflict Zone")
            for i, alert in enumerate(report_data['top_alerts'], 1):
                csv_lines.append(f"{i},{alert['title']},{alert['score']},{alert['source']},{alert['published']},{alert['conflict_group']}")
            
            # Write to file
            with open(report_file, 'w') as f:
                f.write("\n".join(csv_lines))
                
            return report_file
        except Exception as e:
            logger.error(f"CSV report generation failed: {str(e)}")
            return self._generate_json_report(report_data, base_filename)
    
    def _generate_excel_report(self, report_data, base_filename):
        """Generate Excel report"""
        try:
            import openpyxl
            report_file = os.path.join(REPORT_DIR, f"{base_filename}.xlsx")
            
            # Create workbook
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "Intel Report"
            
            # Header
            ws.append(["MILITARY-GLOBAL CONFLICT INTELLIGENCE REPORT"])
            ws.append([f"Generated: {report_data['timestamp']}"])
            ws.append([])
            
            # Threat level
            threat_level = report_data['threat_level']
            ws.append(["Global Threat Level:", threat_level[1]])
            ws.append(["Description:", threat_level[2]])
            ws.append([])
            
            # Daily stats
            ws.append(["Daily Statistics"])
            ws.append(["Metric", "Value"])
            stats = report_data['daily_stats']
            ws.append(["Killed", stats['killed']])
            ws.append(["Injured", stats['injured']])
            ws.append(["Civilian Casualties", stats['civilian_casualties']])
            ws.append(["Strikes/Attacks", stats['strikes']])
            ws.append(["Active Conflicts", stats['conflict_count']])
            ws.append([])
            
            # Conflict stats
            ws.append(["Conflict Zone Analysis"])
            ws.append(["Conflict Zone", "Killed", "Injured", "Incidents"])
            for conflict, killed, injured, reports in report_data['conflict_stats']:
                ws.append([conflict, killed, injured, reports])
            ws.append([])
            
            # Top alerts
            ws.append(["Top Threat Alerts"])
            ws.append(["Rank", "Title", "Threat Score", "Source", "Published", "Conflict Zone"])
            for i, alert in enumerate(report_data['top_alerts'], 1):
                ws.append([i, alert['title'], alert['score'], alert['source'], alert['published'], alert['conflict_group']])
            
            # Save workbook
            wb.save(report_file)
            return report_file
        except Exception as e:
            logger.error(f"Excel report generation failed: {str(e)}")
            return self._generate_csv_report(report_data, base_filename)
            
    def _generate_executive_summary(self, report_data):
        """Generate detailed executive summary"""
        stats = report_data['daily_stats']
        threat_level = report_data['threat_level']
        
        summary = f"Global Conflict Intelligence Report. As of {report_data['timestamp']}, "
        summary += f"the worldwide threat level is {threat_level[1]}. "
        
        if stats:
            summary += f"Today's conflict statistics show {stats['killed']} fatalities, "
            summary += f"{stats['injured']} injuries, and {stats['civilian_casualties']} civilian casualties. "
            summary += f"There were {stats['strikes']} reported strikes across {stats['conflict_count']} active conflict zones. "
        
        if report_data['conflict_stats']:
            deadliest_conflict = max(
                report_data['conflict_stats'], 
                key=lambda x: x[1], 
                default=("", 0, 0, 0)
            )
            if deadliest_conflict[1] > 0:
                summary += f"The most severe conflict is {deadliest_conflict[0]} with {deadliest_conflict[1]} fatalities. "
        
        if report_data['top_alerts']:
            top_alert = report_data['top_alerts'][0]
            summary += f"The highest priority alert is: {top_alert['title']} with a threat score of {top_alert['score']}. "
            summary += f"This was reported by {top_alert['source']}. "
        
        # Add more detail for written reports
        summary += "\n\nDetailed Analysis:\n"
        summary += "The current global security situation remains volatile with multiple hotspots showing increased activity. "
        
        if stats and stats['killed'] > 100:
            summary += "The high casualty figures indicate escalating violence in several regions. "
        
        if threat_level[0] >= 70:
            summary += "Critical threats require immediate attention and contingency planning. "
        
        summary += "Intelligence suggests that the following areas require close monitoring: "
        conflicts = [c[0] for c in report_data['conflict_stats'][:3]]
        summary += ", ".join(conflicts) + ". "
        
        summary += "Recommend increased surveillance and readiness in these zones."
        
        return summary

# ======================
# LIVE DASHBOARD
# ======================
class LiveDashboard:
    def __init__(self, engine):
        self.engine = engine
        self.dashboard_process = None
        self.audio = AudioBriefing()
        self.update_thread = None
        self.update_running = False
        
    def start(self, intel_data=None):
        """Start live dashboard"""
        if not intel_data:
            print("⚠️ No intelligence data. Run Threat Scan first.")
            return
            
        # Create dashboard HTML
        dashboard_file = os.path.join(DASHBOARD_DIR, "live_dashboard.html")
        self._generate_dashboard_html(intel_data, dashboard_file)
        
        # Start dashboard server in background
        if self.dashboard_process:
            self.dashboard_process.terminate()
            
        self.dashboard_process = subprocess.Popen(
            [sys.executable, "-m", "http.server", "8050", "--directory", DASHBOARD_DIR],
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL
        )
        
        # Start live audio updates
        self.start_live_updates(intel_data)
        
        # Open in browser
        webbrowser.open("http://localhost:8050/live_dashboard.html")
        print("\n🌐 Live dashboard running at http://localhost:8050/live_dashboard.html")
        print("Press Ctrl+C in this terminal to stop the dashboard")
        
        return True
        
    def stop(self):
        """Stop live dashboard"""
        if self.dashboard_process:
            self.dashboard_process.terminate()
            self.dashboard_process = None
            
        if self.update_thread:
            self.update_running = False
            self.update_thread.join()
            
        print("\n🛑 Dashboard stopped")
        
    def start_live_updates(self, intel_data):
        """Start live audio updates for dashboard"""
        if self.update_thread and self.update_thread.is_alive():
            self.update_running = False
            self.update_thread.join()
            
        self.update_running = True
        self.update_thread = threading.Thread(target=self._update_worker, args=(intel_data,))
        self.update_thread.daemon = True
        self.update_thread.start()
        
    def _update_worker(self, intel_data):
        """Worker thread for live audio updates"""
        update_interval = 300  # 5 minutes
        while self.update_running:
            # Generate update briefing
            update_text = self.audio.generate_dashboard_briefing(intel_data)
            if update_text:
                audio_file = self.audio.generate_briefing(update_text, "dashboard_update.mp3")
                # Play update
                if audio_file:
                    self.audio.play_briefing(audio_file)
            
            # Wait for next update
            time.sleep(update_interval)
        
    def _generate_dashboard_html(self, intel_data, output_file):
        """Generate interactive dashboard HTML"""
        # Prepare data for visualization
        conflicts = []
        casualties = []
        threat_scores = []
        
        for alert in intel_data:
            conflicts.append(alert['conflict_group'])
            casualties.append(alert.get('casualties_killed', 0))
            threat_scores.append(alert['score'])
        
        # Create HTML with Plotly visualizations
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>MG-CIS Live Dashboard</title>
            <script src="https://cdn.plot.ly/plotly-latest.min.js"></script>
            <style>
                body {{ font-family: Arial, sans-serif; background-color: #0f0f23; color: #e0e0e0; }}
                .dashboard {{ max-width: 1200px; margin: 0 auto; padding: 20px; }}
                .header {{ text-align: center; margin-bottom: 30px; }}
                .header h1 {{ color: #ff6600; font-size: 32px; }}
                .header p {{ font-size: 18px; }}
                .charts {{ display: flex; flex-wrap: wrap; justify-content: space-between; }}
                .chart {{ background-color: #1a1a2e; border-radius: 8px; padding: 15px; margin-bottom: 20px; 
                         box-shadow: 0 4px 8px rgba(0,0,0,0.3); }}
                .full {{ width: 100%; }}
                .half {{ width: 48%; }}
                .controls {{ text-align: center; margin: 20px 0; }}
                button {{ background-color: #ff6600; color: white; border: none; padding: 10px 20px;
                         border-radius: 4px; cursor: pointer; font-size: 16px; margin: 0 10px; }}
                button:hover {{ background-color: #e65c00; }}
                .alert-list {{ max-height: 400px; overflow-y: auto; }}
                .alert-item {{ background-color: #2a2a40; padding: 10px; margin: 10px 0; border-radius: 4px; }}
                .alert-title {{ font-weight: bold; color: #ff6600; }}
                .alert-score {{ color: #ff3300; font-weight: bold; }}
            </style>
        </head>
        <body>
            <div class="dashboard">
                <div class="header">
                    <h1>MG-CIS LIVE DASHBOARD</h1>
                    <p>Real-time Global Conflict Intelligence</p>
                </div>
                
                <div class="controls">
                    <button onclick="window.location.reload()">Refresh Data</button>
                    <button onclick="playAudio()">Play Audio Briefing</button>
                    <button onclick="window.location.href='main_menu.html'">Back to Main</button>
                    <button onclick="window.close()">Exit Dashboard</button>
                </div>
                
                <div class="charts">
                    <div class="chart full">
                        <h2>Global Threat Level</h2>
                        <div id="threatGauge"></div>
                    </div>
                    
                    <div class="chart half">
                        <h2>Conflict Zone Analysis</h2>
                        <div id="conflictChart"></div>
                    </div>
                    
                    <div class="chart half">
                        <h2>Threat Score Distribution</h2>
                        <div id="threatChart"></div>
                    </div>
                    
                    <div class="chart full">
                        <h2>Top Threat Alerts</h2>
                        <div class="alert-list">
        """
        
        # Add alerts
        for i, alert in enumerate(intel_data[:10]):
            html_content += f"""
            <div class="alert-item">
                <div class="alert-title">Alert #{i+1}: {alert['title']}</div>
                <div class="alert-score">Threat Score: {alert['score']}</div>
                <div>Conflict: {alert['conflict_group']}</div>
                <div>Casualties: {alert.get('casualties_killed', 0)} killed, {alert.get('casualties_injured', 0)} injured</div>
                <div>Source: {alert['source']} | Published: {alert['published']}</div>
                <div>Matched Terms: {', '.join(alert['matched'])}</div>
            </div>
            """
        
        html_content += """
                        </div>
                    </div>
                </div>
            </div>
            
            <script>
                // Threat Gauge
                var threatGauge = {{
                    type: "indicator",
                    mode: "gauge+number",
                    value: 0,
                    gauge: {{
                        axis: {{ range: [0, 100] }},
                        bar: {{ color: "darkred" }},
                        steps: [
                            {{ range: [0, 30], color: "green" }},
                            {{ range: [30, 50], color: "yellow" }},
                            {{ range: [50, 70], color: "orange" }},
                            {{ range: [70, 90], color: "red" }},
                            {{ range: [90, 100], color: "darkred" }}
                        ]
                    }}
                }};
                
                // Conflict Chart
                var conflictChart = {{
                    type: 'bar',
                    x: [],
                    y: [],
                    marker: {{ color: 'rgb(255,102,0)' }}
                }};
                
                // Threat Distribution
                var threatChart = {{
                    type: 'histogram',
                    x: [],
                    marker: {{ color: 'rgb(255,102,0)' }}
                }};
                
                // Layouts
                var gaugeLayout = {{
                    title: 'Current Threat Level',
                    font: {{ color: '#e0e0e0' }}
                }};
                
                var conflictLayout = {{
                    title: 'Casualties by Conflict Zone',
                    font: {{ color: '#e0e0e0' }},
                    plot_bgcolor: '#1a1a2e',
                    paper_bgcolor: '#1a1a2e',
                    xaxis: {{ title: 'Conflict Zone', color: '#e0e0e0' }},
                    yaxis: {{ title: 'Casualties', color: '#e0e0e0' }}
                }};
                
                var threatLayout = {{
                    title: 'Alert Threat Scores',
                    font: {{ color: '#e0e0e0' }},
                    plot_bgcolor: '#1a1a2e',
                    paper_bgcolor: '#1a1a2e',
                    xaxis: {{ title: 'Threat Score', range: [0, 100], color: '#e0e0e0' }},
                    yaxis: {{ title: 'Number of Alerts', color: '#e0e0e0' }}
                }};
                
                // Update with actual data
                threatGauge.value = {max(a['score'] for a in intel_data) if intel_data else 0};
                
                conflictChart.x = {json.dumps(conflicts)};
                conflictChart.y = {json.dumps(casualties)};
                
                threatChart.x = {json.dumps(threat_scores)};
                
                // Plot charts
                Plotly.newPlot('threatGauge', [threatGauge], gaugeLayout);
                Plotly.newPlot('conflictChart', [conflictChart], conflictLayout);
                Plotly.newPlot('threatChart', [threatChart], threatLayout);
                
                // Audio playback
                function playAudio() {{
                    var audio = new Audio('/audio_briefings/dashboard_briefing.mp3');
                    audio.play();
                }}
            </script>
        </body>
        </html>
        """
        
        with open(output_file, 'w') as f:
            f.write(html_content)
            
        # Create main menu page
        with open(os.path.join(DASHBOARD_DIR, "main_menu.html"), "w") as f:
            f.write("""
            <!DOCTYPE html>
            <html>
            <head>
                <title>MG-CIS Main Menu</title>
                <style>
                    body { font-family: Arial, sans-serif; background-color: #0f0f23; color: #e0e0e0; text-align: center; }
                    .container { max-width: 600px; margin: 100px auto; }
                    h1 { color: #ff6600; }
                    a { display: block; padding: 15px; margin: 20px; background: #1a1a2e; color: #ff6600; 
                        text-decoration: none; border-radius: 5px; font-size: 20px; }
                    a:hover { background: #2a2a40; }
                </style>
            </head>
            <body>
                <div class="container">
                    <h1>MG-CIS MAIN MENU</h1>
                    <a href="live_dashboard.html">Return to Dashboard</a>
                    <a href="#" onclick="window.close()">Exit System</a>
                </div>
            </body>
            </html>
            """)
            
        # Generate audio briefing
        briefing_text = self.audio.generate_dashboard_briefing(intel_data)
        audio_file = self.audio.generate_briefing(briefing_text, "dashboard_briefing.mp3")
        if audio_file:
            shutil.copy(audio_file, os.path.join(DASHBOARD_DIR, "dashboard_briefing.mp3"))
        
        return True

# ======================
# REAL-TIME ALERT SYSTEM
# ======================
class AlertSystem:
    def __init__(self):
        self.alert_history = deque(maxlen=100)
        self.alert_lock = threading.Lock()
        self.alert_threshold = int(config['DEFAULT'].get('alert_threshold', 80))
        self.audio = AudioBriefing()

    def trigger_alert(self, alert_data):
        """Display high-priority alerts immediately"""
        with self.alert_lock:
            self.alert_history.append(alert_data)
            
            # Create visual alert with flashing effect
            border = "✖" * 100
            alert_title = f"🚨🚨 CRITICAL THREAT ALERT 🚨🚨"
            
            # Prepare alert content
            content = [
                border,
                colored(alert_title, "red", attrs=["bold", "blink"]),
                "",
                colored(f"▸ Threat Level: {alert_data['score']}/100", "red"),
                colored(f"▸ Conflict Zone: {alert_data['conflict_group']}", "yellow"),
                colored(f"▸ Casualties: {alert_data.get('casualties_killed', 0)} killed, {alert_data.get('casualties_injured', 0)} injured", "magenta"),
                colored(f"▸ Source: {alert_data['source']}", "cyan"),
                "",
                colored(alert_data['title'], "white", attrs=["bold"]),
                "",
                colored(f"▸ Matched Terms: {', '.join(alert_data['matched'])}", "magenta"),
                colored(f"▸ Published: {alert_data['published']}", "blue"),
                "",
                colored(f"🔗 {alert_data['link']}", "blue"),
                border
            ]
            
            # Print with flashing effect
            for _ in range(3):  # Flash 3 times for attention
                for line in content:
                    print(line)
                time.sleep(0.5)
                if _ < 2:  # Clear screen between flashes except last
                    print("\033[H\033[J")  # ANSI escape to clear screen

    async def log_alert(self, alert_data):
        """Log alert to database"""
        try:
            async with aiosqlite.connect(DB_FILE) as db:
                await db.execute("""
                INSERT INTO high_priority_alerts (
                    title, url, source, threat_score, conflict_group, matched_terms
                ) VALUES (?, ?, ?, ?, ?, ?)
                """, (
                    alert_data.get("title", "Untitled"),
                    alert_data.get("link", "#"),
                    alert_data.get("source", "Unknown"),
                    alert_data.get("score", 0),
                    alert_data.get("conflict_group", "Global"),
                    json.dumps(alert_data.get("matched", []))
                ))
                await db.commit()
        except Exception as e:
            logger.error(f"Alert logging error: {str(e)}")

    async def process_alerts(self, alerts):
        """Process alerts and trigger high-priority notifications"""
        for alert in alerts:
            if alert['score'] >= self.alert_threshold:
                # Trigger real-time alert
                self.trigger_alert(alert)
                
                # Log to database
                await self.log_alert(alert)
                
                # Add audible alert (system bell)
                print('\a')  # System alert sound
                
                # Generate audio briefing
                if self.audio.enabled:
                    threat_level = self.audio.get_threat_level_name(alert['score'])
                    briefing_text = f"Critical threat alert. {alert['title']} in {alert['conflict_group']}. Threat level: {threat_level}. "
                    if alert.get('casualties_killed', 0) > 0:
                        briefing_text += f"Casualties: {alert['casualties_killed']} killed."
                    audio_file = self.audio.generate_briefing(briefing_text, f"alert_{time.time()}.mp3")
                    if audio_file:
                        self.audio.play_briefing(audio_file)
                
                # Pause briefly for operator attention
                time.sleep(3)

# ======================
# AI-POWERED ANALYSIS
# ======================
class ThreatAnalyzer:
    def __init__(self):
        self.cluster_models = {}
        self.vectorizers = {}
        self.cluster_labels = {}
        self.translator = GoogleTranslator()
        self.llm_enabled = config['DEFAULT'].getboolean('llm_integration', False)
        
    def create_ssl_context(self):
        """Create optimized SSL context"""
        ctx = ssl.create_default_context()
        ctx.set_ciphers('HIGH:!aNULL:!eNULL:!EXPORT:!DES:!RC4:!3DES:!MD5:!PSK')
        if hasattr(ctx, "minimum_version"):
            ctx.minimum_version = ssl.TLSVersion.TLSv1_2
        return ctx

    def get_threat_level(self, score):
        """Get threat level based on score"""
        for threshold, (symbol, name, desc) in sorted(THREAT_LEVELS.items(), reverse=True):
            if score >= threshold:
                return symbol, name, desc
        return THREAT_LEVELS[0]

    def preprocess_text(self, text):
        """Text normalization with stemming and stopword removal"""
        if not stemmer or not stop_words:
            return set()
            
        text = re.sub(r'[^\w\s#]', '', text.lower())
        tokens = word_tokenize(text)
        return {stemmer.stem(word) for word in tokens if word not in stop_words}

    def translate_content(self, text):
        """Translate content to English if needed"""
        try:
            if len(text) > 5000:  # Limit translation to reasonable length
                text = text[:5000]
            return self.translator.translate(text)
        except Exception as e:
            logger.warning(f"Translation error: {str(e)}")
            return text  # Return original if translation fails

    def detect_language(self, text):
        """Detect language of text"""
        try:
            if len(text) > 50:
                return self.translator.detect(text[:500])
            return "en"
        except:
            return "en"

    def calculate_threat_score(self, article):
        """AI-powered threat assessment with weighted scoring"""
        content = f"{article.get('title', '')} {article.get('description', '')}"
        
        # Detect and translate non-English content
        lang = self.detect_language(content)
        if lang != 'en':
            content = self.translate_content(content)
            article['language'] = lang
        
        processed = self.preprocess_text(content)
        
        score = 0
        matched_terms = []
        
        # Keyword scoring with weights
        for category, terms in KEYWORD_GROUPS.items():
            for term, weight in terms.items():
                if term in processed:
                    score += weight
                    matched_terms.append(term)
        
        # Recency bonus
        if 'published_parsed' in article:
            try:
                pub_date = datetime.datetime(*article.published_parsed[:6])
                recency = (datetime.datetime.now() - pub_date).total_seconds() / 3600
                score += max(0, 25 - recency)  # Higher bonus for recent news
            except:
                pass
        
        # Source credibility bonus
        source = getattr(article.feed, 'title', '').lower() if hasattr(article, 'feed') else ''
        if 'reuters' in source or 'ap' in source or 'defense' in source:
            score += 10
        elif 'bbc' in source or 'guardian' in source or 'aljazeera' in source:
            score += 7
            
        # Sentiment analysis
        if sia:
            sentiment = sia.polarity_scores(content)
            if sentiment['compound'] < -0.5:  # Highly negative sentiment
                score += 8
        
        return min(100, score), list(set(matched_terms))

    def cluster_articles(self, articles, n_clusters=5):
        """Cluster articles using K-means for thematic analysis"""
        if not articles or not hasattr(self, 'vectorizers'):
            return articles
            
        # Prepare text data
        texts = [f"{a['title']} {a.get('summary', '')}" for a in articles]
        
        # Create TF-IDF vectors
        if 'tfidf' not in self.vectorizers:
            self.vectorizers['tfidf'] = TfidfVectorizer(
                stop_words='english', 
                max_features=1000
            )
            tfidf_matrix = self.vectorizers['tfidf'].fit_transform(texts)
        else:
            tfidf_matrix = self.vectorizers['tfidf'].transform(texts)
        
        # Cluster using K-means
        if 'kmeans' not in self.cluster_models:
            self.cluster_models['kmeans'] = KMeans(
                n_clusters=min(n_clusters, len(articles)),
                random_state=42
            )
            clusters = self.cluster_models['kmeans'].fit_predict(tfidf_matrix)
        else:
            clusters = self.cluster_models['kmeans'].predict(tfidf_matrix)
        
        # Assign cluster labels to articles
        for i, article in enumerate(articles):
            article['cluster'] = int(clusters[i])
            
        # Generate cluster labels
        self.generate_cluster_labels(articles, tfidf_matrix)
        return articles

    def generate_cluster_labels(self, articles, tfidf_matrix):
        """Generate descriptive labels for clusters"""
        if 'kmeans' not in self.cluster_models:
            return
            
        # Get cluster centers
        cluster_centers = self.cluster_models['kmeans'].cluster_centers_
        terms = self.vectorizers['tfidf'].get_feature_names_out()
        
        self.cluster_labels = {}
        for i, center in enumerate(cluster_centers):
            # Get top terms for this cluster
            top_terms = [
                terms[idx] 
                for idx in center.argsort()[:-6:-1]
            ]
            self.cluster_labels[i] = ", ".join(top_terms)
            # Assign label to articles
            for article in articles:
                if article['cluster'] == i:
                    article['cluster_label'] = self.cluster_labels[i]

    def predict_conflict_escalation(self, conflict_data):
        """Predict conflict escalation using historical patterns"""
        # Simplified prediction model
        escalation_score = 0
        
        # Factors influencing escalation
        factors = {
            'casualties': 0.4,
            'duration': 0.2,
            'international_involvement': 0.3,
            'recent_activity': 0.5
        }
        
        # Calculate escalation score
        escalation_score += min(conflict_data.get('casualties', 0) / 1000, 10) * factors['casualties']
        escalation_score += min(conflict_data.get('duration_months', 0) / 12, 5) * factors['duration']
        escalation_score += conflict_data.get('international_involvement', 0) * factors['international_involvement']
        escalation_score += conflict_data.get('recent_activity', 0) * factors['recent_activity']
        
        # Cap at 100
        escalation_score = min(escalation_score * 10, 100)
        
        return escalation_score

    def generate_tool_code(self, tool_description):
        """AI-generated code for custom OSINT tools"""
        # This is a simplified version - in reality this would connect to an AI API
        tool_name = tool_description.lower().replace(" ", "_")[:20] + "_tool"
        
        code = f"""#!/usr/bin/env python3
# AUTO-GENERATED OSINT TOOL: {tool_description}
# Generated at: {datetime.datetime.now()}

import requests
import json
from bs4 import BeautifulSoup

def run_tool():
    \"\"\"{tool_description}\"\"\"
    print("Running OSINT tool: {tool_description}")
    # TODO: Implement tool functionality
    # Example: result = requests.get("https://api.threatintel.example")
    return {{"status": "success", "data": "Tool executed"}}

if __name__ == "__main__":
    result = run_tool()
    print(json.dumps(result, indent=2))
"""
        return tool_name, code

# ======================
# CUSTOM TOOL MANAGER
# ======================
class ToolManager:
    def __init__(self):
        self.tools_dir = TOOL_DIR
        
    async def create_tool(self, description):
        """Create a new custom OSINT tool"""
        analyzer = ThreatAnalyzer()
        tool_name, code = analyzer.generate_tool_code(description)
        tool_path = os.path.join(self.tools_dir, f"{tool_name}.py")
        
        # Save tool code
        with open(tool_path, "w") as f:
            f.write(code)
        os.chmod(tool_path, 0o700)
        
        # Save to database
        async with aiosqlite.connect(DB_FILE) as db:
            await db.execute("""
            INSERT INTO tools (name, description, code)
            VALUES (?, ?, ?)
            """, (tool_name, description, code))
            await db.commit()
            
        return tool_name, tool_path

    async def list_tools(self):
        """List all available custom tools"""
        tools = []
        async with aiosqlite.connect(DB_FILE) as db:
            cursor = await db.execute("SELECT name, description FROM tools")
            tools = await cursor.fetchall()
        return tools

    def run_tool(self, tool_name):
        """Execute a custom OSINT tool"""
        tool_path = os.path.join(self.tools_dir, f"{tool_name}.py")
        if not os.path.exists(tool_path):
            return {"status": "error", "message": "Tool not found"}
            
        try:
            result = subprocess.run(
                [sys.executable, tool_path],
                capture_output=True,
                text=True,
                timeout=60
            )
            if result.returncode == 0:
                return {"status": "success", "output": result.stdout}
            else:
                return {"status": "error", "message": result.stderr}
        except Exception as e:
            return {"status": "error", "message": str(e)}

# ======================
# OSINT COLLECTION ENGINE
# ======================
class IntelEngine:
    def __init__(self):
        self.session = aiohttp.ClientSession()
        # Improved casualty patterns with word boundaries
        self.casualty_patterns = {
            'killed': re.compile(r'\b(\d+)\s*(?:killed|dead|fatalit(y|ies)|death(s)?|died)\b', re.IGNORECASE),
            'injured': re.compile(r'\b(\d+)\s*(?:injured|wounded|hurt)\b', re.IGNORECASE),
            'civilian': re.compile(r'\b(\d+)\s*(?:civilian\s+casualt(y|ies)|civilians\s+killed|non-combatants)\b', re.IGNORECASE),
            'strike': re.compile(r'\b(\d+)\s*(?:strike|attack|raid|bombing)\b', re.IGNORECASE)
        }
        max_requests = int(config['DEFAULT'].get('max_concurrent_requests', 10))
        self.limiter = asyncio.Semaphore(max_requests)
        self.analyzer = ThreatAnalyzer()
        self.tool_manager = ToolManager()
        self.alert_system = AlertSystem()
        self.geopy_enabled = config['DEFAULT'].getboolean('geospatial_enabled', True)
        self.social_media_enabled = config['DEFAULT'].getboolean('social_media_enabled', True)
        self.report_gen = ReportGenerator()
        self.dashboard = LiveDashboard(self)
        self.feed_cache = {}
        self.cache_file = os.path.join(DB_DIR, "feed_cache.json")
        self.load_cache()

    def load_cache(self):
        """Load feed cache from file"""
        if os.path.exists(self.cache_file):
            try:
                with open(self.cache_file, 'r') as f:
                    self.feed_cache = json.load(f)
            except:
                self.feed_cache = {}

    def save_cache(self):
        """Save feed cache to file"""
        try:
            with open(self.cache_file, 'w') as f:
                json.dump(self.feed_cache, f)
        except:
            pass

    async def __aenter__(self):
        return self

    async def __aexit__(self, exc_type, exc, tb):
        await self.session.close()
        self.save_cache()

    async def find_rss_feeds(self, site_url):
        """Discover RSS feeds from a website with enhanced discovery"""
        try:
            async with self.session.get(site_url, timeout=10, ssl=self.analyzer.create_ssl_context()) as response:
                html_content = await response.text()
                soup = BeautifulSoup(html_content, 'html.parser')
                
            feeds = []
            # Check for RSS link tags
            for link in soup.find_all('link', type=lambda t: t and 'rss' in t):
                href = link.get('href')
                if href:
                    if href.startswith('/'):
                        feeds.append(f"{urlparse(site_url).scheme}://{urlparse(site_url).netloc}{href}")
                    else:
                        feeds.append(href)
            
            # Check for common RSS paths
            common_paths = [
                "/rss", "/feed", "/rss.xml", "/feed.xml", "/atom.xml",
                "/rss/all", "/rss/news", "/rss/world", "/news/rss", "/world/rss"
            ]
            base_url = f"{urlparse(site_url).scheme}://{urlparse(site_url).netloc}"
            for path in common_paths:
                feeds.append(f"{base_url}{path}")
            
            # Check for hyperlinks containing "rss" or "feed"
            for a_tag in soup.find_all('a', href=True):
                href = a_tag['href']
                if any(keyword in href.lower() for keyword in ['rss', 'feed', 'xml']):
                    if href.startswith('/'):
                        feeds.append(f"{base_url}{href}")
                    elif href.startswith('http'):
                        feeds.append(href)
            
            return list(set(feeds))
        except Exception as e:
            logger.warning(f"Discovery error for {site_url}: {str(e)}")
            return []

    async def discover_feeds_from_website_list(self):
        """Discover RSS feeds from all provided websites"""
        print("\n🔍 Discovering RSS feeds from global news sources...")
        tasks = [self.find_rss_feeds(url) for url in WEBSITE_LIST]
        results = await asyncio.gather(*tasks)
        
        valid_feeds = []
        for feed_list in results:
            valid_feeds.extend(feed_list)
        
        # Save discovered feeds
        self.save_feeds(valid_feeds)
        return valid_feeds

    def save_feeds(self, feeds):
        """Save feeds to encrypted storage"""
        try:
            data = json.dumps(feeds).encode()
            encrypted = cipher_suite.encrypt(data)
            with open(FEED_STORAGE, "wb") as f:
                f.write(encrypted)
        except Exception as e:
            logger.error(f"Failed to save feeds: {str(e)}")
            # Attempt to save unencrypted as fallback
            try:
                with open(FEED_STORAGE + ".bak", "w") as f:
                    json.dump(feeds, f)
                logger.warning("Saved feeds in unencrypted backup")
            except:
                logger.error("Failed to save backup feeds")

    def load_feeds(self):
        """Load feeds from encrypted storage with decryption fallback"""
        if not os.path.exists(FEED_STORAGE):
            return []
            
        try:
            with open(FEED_STORAGE, "rb") as f:
                encrypted = f.read()
                decrypted = cipher_suite.decrypt(encrypted)
                return json.loads(decrypted.decode())
        except Exception as e:
            logger.error(f"Feed decryption failed: {str(e)}")
            # Attempt to load from unencrypted backup
            try:
                if os.path.exists(FEED_STORAGE + ".bak"):
                    with open(FEED_STORAGE + ".bak", "r") as f:
                        return json.load(f)
            except:
                logger.error("Failed to load backup feeds")
            return []

    async def fetch_feed(self, url):
        """Fetch and parse RSS feed with retry mechanism and caching"""
        # Check cache first
        cache_key = hashlib.md5(url.encode()).hexdigest()
        now = time.time()
        
        if cache_key in self.feed_cache:
            cache_data = self.feed_cache[cache_key]
            # Use cache if it's less than 30 minutes old
            if now - cache_data['timestamp'] < 1800:
                return feedparser.parse(cache_data['content'])
        
        try:
            ssl_context = self.analyzer.create_ssl_context()
            async with self.limiter:
                headers = {}
                if cache_key in self.feed_cache:
                    cache_data = self.feed_cache[cache_key]
                    headers['If-None-Match'] = cache_data.get('etag', '')
                    headers['If-Modified-Since'] = cache_data.get('last_modified', '')
                
                async with self.session.get(url, ssl=ssl_context, timeout=15, headers=headers) as response:
                    if response.status == 304:
                        # Not modified - use cache
                        return feedparser.parse(self.feed_cache[cache_key]['content'])
                    
                    content = await response.text()
                    # Update cache
                    self.feed_cache[cache_key] = {
                        'content': content,
                        'timestamp': now,
                        'etag': response.headers.get('ETag', ''),
                        'last_modified': response.headers.get('Last-Modified', '')
                    }
                    return feedparser.parse(content)
        except Exception as e:
            logger.warning(f"Error fetching {url}: {str(e)}")
            return None

    async def analyze_global_threats(self, disable_audio=False):
        """Main intelligence processing pipeline with performance optimizations"""
        feeds = self.load_feeds()
        if not feeds:
            print("⚠️ No feeds available. Discovering new sources...")
            feeds = await self.discover_feeds_from_website_list()
            
        # Disable audio during scan if requested
        original_audio_setting = self.alert_system.audio.enabled
        if disable_audio:
            self.alert_system.audio.enabled = False
            
        tasks = [self.fetch_feed(url) for url in feeds]
        results = await asyncio.gather(*tasks)
        
        alerts = []
        # Process entries in parallel
        entry_tasks = []
        for feed in results:
            if not feed or not getattr(feed, 'entries', None):
                continue
                
            for entry in feed.entries:
                entry_tasks.append(self.process_entry(feed, entry))
                
        # Process entries concurrently
        processed_entries = await asyncio.gather(*entry_tasks)
        alerts = [entry for entry in processed_entries if entry is not None]
        
        # Apply clustering
        alerts = self.analyzer.cluster_articles(alerts)
        
        # Process through real-time alert system
        await self.alert_system.process_alerts(alerts)
        
        # Restore audio setting
        if disable_audio:
            self.alert_system.audio.enabled = original_audio_setting
            
        return sorted(alerts, key=lambda x: x['score'], reverse=True)
    
    async def process_entry(self, feed, entry):
        """Process a single feed entry asynchronously"""
        try:
            score, matched = self.analyzer.calculate_threat_score(entry)
            if score > 25:  # Threshold for important news
                # Extract conflict group
                conflict_group = "Global"
                for conflict, pattern in CONFLICT_GROUPS.items():
                    if pattern.search(entry.get('title', '') + entry.get('description', '')):
                        conflict_group = conflict
                        break
                
                # Extract casualties with improved patterns
                casualties = self.extract_casualties(entry.get('description', ''))
                
                alert = {
                    'title': entry.get('title', 'No title'),
                    'link': entry.get('link', '#'),
                    'score': score,
                    'matched': matched,
                    'source': getattr(feed.feed, 'title', 'Unknown'),
                    'published': entry.get('published', 'Unknown date'),
                    'summary': entry.get('description', ''),
                    'conflict_group': conflict_group,
                    'casualties_killed': casualties['killed'],
                    'casualties_injured': casualties['injured'],
                    'casualties_civilian': casualties['civilian']
                }
                
                # Save to database with proper encryption handling
                await self.save_article({
                    'title': entry.get('title', ''),
                    'url': entry.get('link', ''),
                    'source': getattr(feed.feed, 'title', 'Unknown'),
                    'published_date': entry.get('published', ''),
                    'summary': entry.get('description', ''),
                    'threat_score': score,
                    'matched_terms': json.dumps(matched),
                    'conflict_group': conflict_group,
                    'casualties_killed': casualties['killed'],
                    'casualties_injured': casualties['injured'],
                    'casualties_civilian': casualties['civilian']
                })
                
                return alert
        except Exception as e:
            logger.error(f"Processing error: {str(e)}")
            return None

    def extract_casualties(self, text):
        """Enhanced casualty extraction with regex patterns"""
        results = {'killed': 0, 'injured': 0, 'civilian': 0, 'strike': 0}
        
        for key, pattern in self.casualty_patterns.items():
            matches = pattern.findall(text)
            if matches:
                try:
                    # Extract numbers from match groups
                    numbers = [int(m[0]) for m in matches if m[0].isdigit()]
                    total = sum(numbers)
                    results[key] = min(total, 10000)  # Prevent unrealistic numbers
                except ValueError:
                    pass
        return results

    async def scrape_article(self, url):
        """Scrape article content from URL with caching"""
        if url in self.scrape_cache:
            return self.scrape_cache[url]
            
        try:
            async with self.limiter:
                async with self.session.get(url, timeout=15, ssl=self.analyzer.create_ssl_context()) as response:
                    html = await response.text()
                    soup = BeautifulSoup(html, "lxml")
                    
                    # Remove unwanted elements
                    for element in soup(["script", "style", "noscript", "iframe", "comment"]):
                        element.decompose()
                    
                    # Extract main content
                    content = ""
                    potential_content = soup.find_all(['article', 'div', 'section'])
                    for elem in potential_content:
                        classes = ' '.join(elem.get('class', [])).lower()
                        elem_id = elem.get('id', '').lower()
                        if any(keyword in classes for keyword in ['article', 'content', 'main', 'story']) or \
                           any(keyword in elem_id for keyword in ['article', 'content']):
                            content = elem.get_text(separator='\n', strip=True)
                            break
                    
                    # Fallback to all paragraphs if no main content found
                    if not content:
                        content = '\n'.join(p.get_text(strip=True) for p in soup.find_all('p'))
                    
                    # Cache the result
                    self.scrape_cache[url] = content
                    return content
        except Exception as e:
            logger.warning(f"Error scraping article: {str(e)}")
            return ""

    async def save_article(self, article):
        """Save article to database with proper encryption handling"""
        try:
            # Encrypt sensitive fields and encode for storage
            encrypted_title = base64.b64encode(
                cipher_suite.encrypt(article.get("title", "").encode())
            ).decode()
            encrypted_summary = base64.b64encode(
                cipher_suite.encrypt(article.get("summary", "").encode())
            ).decode()
            
            async with aiosqlite.connect(DB_FILE) as db:
                await db.execute("""
                INSERT OR IGNORE INTO articles (
                    title, url, source, published_date, summary, 
                    threat_score, matched_terms, conflict_group, language,
                    casualties_killed, casualties_injured, casualties_civilian
                ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, (
                    encrypted_title,
                    article.get("url", "#"),
                    article.get("source", "Unknown"),
                    article.get("published_date", datetime.datetime.now().isoformat()),
                    encrypted_summary,
                    article.get("threat_score", 0),
                    article.get("matched_terms", "[]"),
                    article.get("conflict_group", "Global"),
                    article.get("language", "en"),
                    article.get("casualties_killed", 0),
                    article.get("casualties_injured", 0),
                    article.get("casualties_civilian", 0)
                ))
                await db.commit()
                return 1
        except Exception as e:
            logger.error(f"Database error: {str(e)}")
            return 0

    async def get_daily_stats(self):
        """Get today's casualty statistics"""
        try:
            async with aiosqlite.connect(DB_FILE) as db:
                today = datetime.date.today().isoformat()
                cursor = await db.execute("""
                SELECT 
                    SUM(casualties_killed) as killed,
                    SUM(casualties_injured) as injured,
                    (SELECT COUNT(*) FROM articles WHERE DATE(timestamp) = DATE('now')) as strikes,
                    SUM(casualties_civilian) as civilian_casualties,
                    COUNT(DISTINCT conflict_group) as conflict_count
                FROM articles
                WHERE DATE(timestamp) = DATE('now')
                """)
                stats = await cursor.fetchone()
                
                if stats:
                    return {
                        "date": today,
                        "killed": stats[0] or 0,
                        "injured": stats[1] or 0,
                        "strikes": stats[2] or 0,
                        "civilian_casualties": stats[3] or 0,
                        "conflict_count": stats[4] or 0
                    }
                return None
        except Exception as e:
            logger.error(f"Database error: {str(e)}")
            return None

    async def get_conflict_stats(self):
        """Get statistics by conflict group"""
        try:
            async with aiosqlite.connect(DB_FILE) as db:
                cursor = await db.execute("""
                SELECT conflict_group, 
                       SUM(casualties_killed) as killed, 
                       SUM(casualties_injured) as injured,
                       COUNT(*) as articles
                FROM articles
                WHERE DATE(timestamp) = DATE('now')
                GROUP BY conflict_group
                ORDER BY killed DESC
                """)
                return await cursor.fetchall()
        except Exception as e:
            logger.error(f"Database error: {str(e)}")
            return []

    async def generate_geospatial_intelligence(self, alerts):
        """Generate geospatial intelligence map"""
        if not alerts or not hasattr(folium, 'Map'):
            return None
            
        # Create base map
        m = folium.Map(location=[30, 0], zoom_start=2)
        
        # Conflict zone coordinates (simplified)
        conflict_zones = {
            "Ukraine": [49, 31],
            "Israel": [31.5, 34.5],
            "Taiwan": [23.5, 121],
            "Kashmir": [34, 76],
            "Syria": [35, 38],
            "Afghanistan": [34, 65],
            "Yemen": [15.5, 48],
            "Ethiopia": [9, 40],
            "Sudan": [15, 30],
            "Myanmar": [22, 96]
        }
        
        # Use geopy for precise geocoding if enabled
        if self.geopy_enabled and hasattr(self, 'geolocator'):
            try:
                for conflict, coords in conflict_zones.items():
                    location = self.geolocator.geocode(conflict)
                    if location:
                        conflict_zones[conflict] = [location.latitude, location.longitude]
            except Exception as e:
                logger.warning(f"Geocoding error: {str(e)}")
        
        # Add markers for each alert
        for alert in alerts[:100]:  # Limit to 100 markers
            # Find matching conflict zone
            for conflict, coords in conflict_zones.items():
                if conflict.lower() in alert['title'].lower() or conflict.lower() in alert['conflict_group'].lower():
                    # Add marker with popup
                    popup_content = f"""
                    <b>{alert['title']}</b><br>
                    Score: {alert['score']}<br>
                    Killed: {alert.get('casualties_killed', 0)}<br>
                    Injured: {alert.get('casualties_injured', 0)}<br>
                    <a href="{alert['link']}" target="_blank">Source</a>
                    """
                    folium.Marker(
                        location=coords,
                        popup=popup_content,
                        icon=folium.Icon(
                            color='red' if alert['score'] > 70 else 'orange',
                            icon='info-sign'
                        )
                    ).add_to(m)
                    break
            else:
                # Add to global conflict marker if no specific location
                folium.Marker(
                    location=[random.uniform(-60, 70), random.uniform(-180, 180)],
                    popup=f"<b>{alert['title']}</b><br>Score: {alert['score']}",
                    icon=folium.Icon(color='blue', icon='info-sign')
                ).add_to(m)
        
        # Save map
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        map_file = os.path.join(HTML_DIR, f"geoint_{timestamp}.html")
        m.save(map_file)
        return map_file

    async def generate_threat_timeline(self, alerts):
        """Generate threat timeline visualization"""
        if not alerts or not hasattr(plt, 'plot'):
            return None
            
        # Prepare data
        timeline_data = []
        for alert in alerts:
            try:
                # Use current date minus random days for simulation
                pub_date = datetime.datetime.now() - datetime.timedelta(days=random.randint(0, 7))
                timeline_data.append({
                    'date': pub_date,
                    'score': alert['score'],
                    'title': alert['title'][:50] + '...' if len(alert['title']) > 50 else alert['title'],
                    'casualties': alert.get('casualties_killed', 0)
                })
            except:
                continue
        
        if not timeline_data:
            return None
            
        # Create DataFrame
        df = pd.DataFrame(timeline_data)
        df['date'] = pd.to_datetime(df['date'])
        df.set_index('date', inplace=True)
        
        # Resample by day
        daily = df['score'].resample('D').max().fillna(0)
        
        # Plot
        plt.figure(figsize=(12, 6))
        plt.plot(daily.index, daily.values, marker='o', linestyle='-', color='darkred')
        plt.fill_between(daily.index, daily.values, color='red', alpha=0.3)
        plt.title('Threat Level Timeline', fontsize=14)
        plt.ylabel('Threat Score', fontsize=12)
        plt.xlabel('Date', fontsize=12)
        plt.grid(True, linestyle='--', alpha=0.7)
        
        # Add casualty markers with error handling
        if not daily.empty:
            max_day = daily.idxmax()
            max_value = daily.max()
            
            # Get total casualties for the day
            try:
                total_casualties = df.loc[df.index.date == max_day.date()]['casualties'].sum()
                plt.annotate(f'{int(total_casualties)} casualties',
                             xy=(max_day, max_value), xytext=(max_day, max_value+5),
                             arrowprops=dict(facecolor='black', shrink=0.05))
            except Exception as e:
                logger.warning(f"Error adding casualty marker: {str(e)}")
        
        # Save plot
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        plot_file = os.path.join(REPORT_DIR, f"threat_timeline_{timestamp}.png")
        plt.savefig(plot_file, bbox_inches='tight')
        plt.close()
        
        return plot_file

    async def monitor_social_media(self):
        """Monitor social media for military intelligence"""
        if not self.social_media_enabled:
            logger.info("Social media monitoring disabled by configuration")
            return []
            
        print("\n📱 Monitoring social media for military intelligence...")
        social_data = []
        
        # Twitter integration with updated v2 API
        twitter_keys = [
            config['DEFAULT'].get('twitter_api_key', ''),
            config['DEFAULT'].get('twitter_api_secret', ''),
            config['DEFAULT'].get('twitter_access_token', ''),
            config['DEFAULT'].get('twitter_access_secret', ''),
            config['DEFAULT'].get('twitter_bearer_token', '')
        ]
        
        # Use Bearer token if available, else use consumer keys
        if twitter_keys[4] and hasattr(sys.modules[__name__], 'tweepy'):  # Bearer token method
            try:
                import tweepy
                client = tweepy.Client(
                    bearer_token=config['DEFAULT']['twitter_bearer_token']
                )
                
                # Search for military conflict keywords
                query = " OR ".join(KEYWORDS[:5] + HASHTAGS[:5])  # Limit to first 5 keywords/hashtags
                tweets = client.search_recent_tweets(
                    query=query,
                    max_results=50,
                    tweet_fields=['created_at', 'author_id'],
                    user_fields=['username'],
                    expansions=['author_id']
                )
                
                # Create user ID to username mapping
                users = {u.id: u.username for u in tweets.includes.get('users', [])} if tweets.includes else {}
                
                for tweet in tweets.data or []:
                    content = tweet.text
                    sentiment = sia.polarity_scores(content)['compound'] if sia else 0
                    
                    conflict_group = "Global"
                    for conflict, pattern in CONFLICT_GROUPS.items():
                        if pattern.search(content):
                            conflict_group = conflict
                            break
                    
                    # Keyword matching
                    matched_keywords = []
                    for keyword in KEYWORDS:
                        if keyword.lower() in content.lower():
                            matched_keywords.append(keyword)
                    
                    social_data.append({
                        "platform": "Twitter",
                        "author": users.get(tweet.author_id, 'unknown'),
                        "content": content,
                        "url": f"https://twitter.com/user/status/{tweet.id}",
                        "post_timestamp": tweet.created_at.isoformat(),
                        "sentiment": sentiment,
                        "conflict_group": conflict_group,
                        "keywords": ", ".join(matched_keywords)
                    })
                    
                logger.info(f"Collected {len(tweets.data or [])} tweets")
            except Exception as e:
                logger.error(f"Twitter monitoring error: {str(e)}")
        elif all(twitter_keys[:4]) and hasattr(sys.modules[__name__], 'tweepy'):  # Consumer key method
            try:
                import tweepy
                auth = tweepy.OAuthHandler(
                    config['DEFAULT']['twitter_api_key'],
                    config['DEFAULT']['twitter_api_secret']
                )
                auth.set_access_token(
                    config['DEFAULT']['twitter_access_token'],
                    config['DEFAULT']['twitter_access_secret']
                )
                api = tweepy.API(auth)
                
                # Search for military conflict keywords
                query = " OR ".join(KEYWORDS[:5] + HASHTAGS[:5])  # Limit to first 5 keywords/hashtags
                tweets = api.search_tweets(q=query, count=50, tweet_mode='extended')
                
                for tweet in tweets:
                    content = tweet.full_text
                    sentiment = sia.polarity_scores(content)['compound'] if sia else 0
                    
                    conflict_group = "Global"
                    for conflict, pattern in CONFLICT_GROUPS.items():
                        if pattern.search(content):
                            conflict_group = conflict
                            break
                    
                    # Keyword matching
                    matched_keywords = []
                    for keyword in KEYWORDS:
                        if keyword.lower() in content.lower():
                            matched_keywords.append(keyword)
                    
                    social_data.append({
                        "platform": "Twitter",
                        "author": tweet.user.screen_name,
                        "content": content,
                        "url": f"https://twitter.com/{tweet.user.screen_name}/status/{tweet.id}",
                        "post_timestamp": tweet.created_at.isoformat(),  # Convert to ISO string
                        "sentiment": sentiment,
                        "conflict_group": conflict_group,
                        "keywords": ", ".join(matched_keywords)
                    })
                    
                logger.info(f"Collected {len(tweets)} tweets")
            except Exception as e:
                logger.error(f"Twitter monitoring error: {str(e)}")
        else:
            # Simulated data collection as fallback
            platforms = ["Twitter", "Telegram", "Reddit", "Facebook", "VKontakte"]
            
            # Generate simulated posts
            for _ in range(20):
                platform = random.choice(platforms)
                author = f"user_{random.randint(10000, 99999)}"
                content = random.choice([
                    f"{random.choice(KEYWORDS)}: {random.choice(HASHTAGS)}",
                    f"Breaking: {random.choice(['Military', 'Government'])} {random.choice(['announcement', 'statement'])}",
                    f"{random.choice(HASHTAGS)} {random.choice(['escalation', 'ceasefire', 'attack'])} reported",
                    f"{random.choice(KEYWORDS)} in {random.choice(['Ukraine', 'Israel', 'Taiwan'])}"
                ])
                url = f"https://{platform}.com/post/{random.randint(100000, 999999)}"
                post_timestamp = datetime.datetime.now() - datetime.timedelta(minutes=random.randint(0, 1440))
                
                # Analyze sentiment
                sentiment = sia.polarity_scores(content)['compound'] if sia else 0
                
                # Determine conflict group
                conflict_group = "Global"
                for conflict, pattern in CONFLICT_GROUPS.items():
                    if pattern.search(content):
                        conflict_group = conflict
                        break
                
                # Keyword matching
                matched_keywords = []
                for keyword in KEYWORDS:
                    if keyword.lower() in content.lower():
                        matched_keywords.append(keyword)
                
                social_data.append({
                    "platform": platform,
                    "author": author,
                    "content": content,
                    "url": url,
                    "post_timestamp": post_timestamp.isoformat(),  # Convert to ISO string
                    "sentiment": sentiment,
                    "conflict_group": conflict_group,
                    "keywords": ", ".join(matched_keywords)
                })
        
        # Save to database
        try:
            async with aiosqlite.connect(DB_FILE) as db:
                for post in social_data:
                    await db.execute("""
                    INSERT INTO social_media (
                        platform, author, content, url, post_timestamp, sentiment, conflict_group, keywords
                    ) VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                    """, (
                        post["platform"],
                        post["author"],
                        post["content"],
                        post["url"],
                        post["post_timestamp"],
                        post["sentiment"],
                        post["conflict_group"],
                        post["keywords"]
                    ))
                await db.commit()
        except Exception as e:
            logger.error(f"Social media save error: {str(e)}")
        
        return social_data

    async def generate_detailed_report(self, alerts, stats):
        """Generate detailed intelligence report"""
        if not alerts:
            return None
            
        # Prepare report data with robust async handling
        try:
            # Gather all async data sources with timeout
            social_media_data = await asyncio.wait_for(
                self.monitor_social_media(), 
                timeout=30
            )
        except (asyncio.TimeoutError, Exception) as e:
            logger.error(f"Social media monitoring failed: {str(e)}")
            social_media_data = []
            
        try:
            conflict_stats = await asyncio.wait_for(
                self.get_conflict_stats(), 
                timeout=15
            )
        except (asyncio.TimeoutError, Exception) as e:
            logger.error(f"Conflict stats failed: {str(e)}")
            conflict_stats = []

        report_data = {
            "timestamp": datetime.datetime.now().isoformat(),
            "threat_level": self.analyzer.get_threat_level(max(a['score'] for a in alerts) if alerts else (0, "Unknown", "No data")),
            "alert_count": len(alerts),
            "top_alerts": sorted(alerts, key=lambda x: x['score'], reverse=True)[:10] if alerts else [],
            "daily_stats": stats,
            "conflict_stats": conflict_stats,
            "social_media": social_media_data[:5]  # Only take top 5 after successful retrieval
        }
        
        return report_data

# ======================
# USER INTERFACE
# ======================
def display_threats(alerts):
    """Display results in terminal with color coding"""
    if not alerts:
        print("\n🟢 No critical alerts detected")
        return
        
    # Calculate global threat level
    max_score = max(a['score'] for a in alerts) if alerts else 0
    symbol, name, desc = THREAT_LEVELS.get(max_score, ("", "", ""))
    
    print(f"\n{symbol} GLOBAL THREAT LEVEL: {name} - {desc}")
    print(f"🚨 GLOBAL THREAT ALERTS ({len(alerts)} detected)")
    print("=" * 100)
    
    # Group alerts by conflict
    conflict_groups = defaultdict(list)
    for alert in alerts:
        conflict_groups[alert.get('conflict_group', 'Global')].append(alert)
    
    # Sort by conflict severity
    for conflict, group in sorted(conflict_groups.items(), key=lambda x: max(a['score'] for a in x[1]), reverse=True):
        print(f"\n🔥 CONFLICT: {conflict} ({len(group)} alerts)")
        print("-" * 100)
        for alert in group[:5]:  # Show top 5 per conflict
            if alert['score'] >= 75:
                color = "red"
            elif alert['score'] >= 50:
                color = "yellow"
            else:
                color = "cyan"
                
            print(colored(f"▓▓ {alert['title']}", color))
            print(f"   🔗 {alert['link']}")
            print(colored(f"   ⚠️  Threat Level: {alert['score']}/100", color))
            print(f"   💀 Killed: {alert.get('casualties_killed', 0)} | 🤕 Injured: {alert.get('casualties_injured', 0)}")
            print(f"   🔍 Matched Terms: {', '.join(alert['matched'])}")
            if 'cluster_label' in alert:
                print(f"   🧩 Threat Cluster: {alert['cluster_label']}")
            print(f"   📰 Source: {alert['source']} | 🕒 {alert['published']}")
            print("-" * 100)

def display_daily_stats(stats):
    """Display daily casualty statistics"""
    if not stats:
        print("\n⚠️ No statistics available for today")
        return
        
    print("\n📊 DAILY CONFLICT STATISTICS")
    print("="*50)
    print(f"📅 Date: {stats['date']}")
    print(f"💀 Killed: {stats['killed']}")
    print(f"🤕 Injured: {stats['injured']}")
    print(f"💥 Strikes/Attacks: {stats['strikes']}")
    print(f"👥 Civilian Casualties: {stats['civilian_casualties']}")
    print(f"🌍 Active Conflicts: {stats['conflict_count']}")
    print("="*50)

def display_conflict_stats(conflict_stats):
    """Display statistics by conflict group"""
    if not conflict_stats:
        print("\n⚠️ No conflict statistics available")
        return
        
    print("\n🔥 CONFLICT ZONE STATISTICS")
    print("=" * 70)
    print(f"{'Conflict':<25} | {'Killed':<10} | {'Injured':<10} | {'Reports':<10}")
    print("-" * 70)
    for conflict, killed, injured, reports in conflict_stats:
        print(f"{conflict:<25} | {killed:<10} | {injured:<10} | {reports:<10}")
    print("=" * 70)

def generate_ascii_banner():
    """Generate ASCII art banner with color"""
    try:
        banner_text = pyfiglet.figlet_format("DAY - R", font="slant")
    except:
        banner_text = "DAY - R\n"
    
    colors = ['red', 'green', 'yellow', 'blue', 'magenta', 'cyan']
    color = random.choice(colors)
    banner = colored(banner_text, color)
    banner += colored("💀 V3N0M0U5  KR0N05 💀\n", 'white', attrs=['bold'])
    banner += colored(" \n")
    banner += colored("Global Conflicts & Threats Detection & Analysis Intelligence System\n", 'white', attrs=['bold'])
    banner += colored(" \n")
    banner += colored("monster666@f5.si"  "\n", 'white', attrs=['bold'])
    banner += colored(" \n")
    banner += colored("💀" * 18 + "\n", 'white', attrs=['bold'])
    return banner

def display_main_menu():
    """Display main command center menu"""
   
#    print("\n" + "⚝"*50)
#    print(generate_ascii_banner())
    print("⚝"*50)
    print(" "*18)
    print("1. Run Global Threat Scan")
    print("2. Discover New Intelligence Sources")
    print("3. View Daily Statistics")
    print("4. View Conflict Zone Statistics")
    print("5. Generate Intelligence Report")
    print("6. Geospatial Intelligence")
    print("7. Threat Timeline Visualization")
    print("8. Monitor Social Media")
    print("9. Historical Conflict Data")
    print("10. Create Custom OSINT Tool")
    print("11. Run Custom OSINT Tool")
    print("12. System Status")
    print("13. Security Audit")
    print("14. View High-Priority Alerts")
    print("15. Configuration Settings")
    print("16. Select Audio Voice")
    print("17. Launch Live Dashboard")
    print("18. Audio Briefings")
    print("0. Exit")
    return input("\n[COMMAND] Select operation: ")

async def view_custom_tools(tool_manager):
    """View and manage custom OSINT tools"""
    while True:
        tools = await tool_manager.list_tools()
        print("\n[INTEL TOOLS]")
        if not tools:
            print("No custom tools available")
        else:
            for i, (name, desc) in enumerate(tools, 1):
                print(f"{i}. {name} - {desc}")
                
        print("\n1. Create New Tool")
        print("2. Execute Tool")
        print("0. Back to Main Menu")
        choice = input("\n[TOOL COMMAND] Select option: ")
        
        if choice == '1':
            description = input("Enter tool description: ")
            tool_name, tool_path = await tool_manager.create_tool(description)
            print(f"✅ Created tool: {tool_name} at {tool_path}")
            
        elif choice == '2' and tools:
            tool_idx = int(input("Select tool to execute: ")) - 1
            if 0 <= tool_idx < len(tools):
                tool_name = tools[tool_idx][0]
                result = tool_manager.run_tool(tool_name)
                print("\n[TOOL OUTPUT]")
                print(json.dumps(result, indent=2))
            else:
                print("❌ Invalid selection")
                
        elif choice == '0':
            return
            
        else:
            print("❌ Invalid selection")

async def view_geospatial_intelligence(engine, alerts):
    """Display geospatial intelligence"""
    if not alerts:
        print("\n⚠️ No intelligence data. Run Threat Scan first.")
        return
        
    print("\n🛰️ Generating geospatial intelligence...")
    map_file = await engine.generate_geospatial_intelligence(alerts)
    if map_file:
        print(f"✅ Geospatial intelligence map saved to: {map_file}")
        webbrowser.open('file://' + os.path.abspath(map_file))
    else:
        print("⚠️ Failed to generate geospatial intelligence")

async def view_threat_timeline(engine, alerts):
    """Display threat timeline visualization"""
    if not alerts:
        print("\n⚠️ No intelligence data. Run Threat Scan first.")
        return
        
    print("\n📈 Generating threat timeline...")
    timeline_file = await engine.generate_threat_timeline(alerts)
    if timeline_file:
        print(f"✅ Threat timeline saved to: {timeline_file}")
        if platform.system() == 'Darwin':
            subprocess.run(['open', timeline_file])
        elif platform.system() == 'Windows':
            os.startfile(timeline_file)
        else:
            subprocess.run(['xdg-open', timeline_file])
    else:
        print("⚠️ Failed to generate threat timeline")

async def view_historical_conflicts():
    """Display historical conflict data"""
    print("\n📜 HISTORICAL CONFLICTS (Last 100 Years)")
    print("=" * 80)
    print(f"{'Conflict':<25} | {'Start':<10} | {'Casualties':<15} | {'Status':<15}")
    print("-" * 80)
    for conflict, data in HISTORICAL_CONFLICTS.items():
        print(f"{conflict:<25} | {data['start']:<10} | {data['casualties']:<15} | {data['status']:<15}")
    print("=" * 80)

def security_audit():
    """Perform security audit of the system"""
    print("\n🔒 SECURITY AUDIT RESULTS")
    print("="*50)
    
    # Check critical security settings
    issues = []
    
    # Check encryption key
    if config['DEFAULT'].get('encryption_key', '') == '':
        issues.append("🔴 Encryption key not configured")
    
    # Check directory permissions
    for directory in [VENV_DIR, REPORT_DIR, DB_DIR, HTML_DIR, TOOL_DIR, CONFIG_DIR, LOG_DIR]:
        if oct(os.stat(directory).st_mode & 0o777) != '0o700':
            issues.append(f"🟠 Incorrect permissions for {directory} - should be 700")
    
    # Check config file permissions
    if os.path.exists(CONFIG_FILE):
        if oct(os.stat(CONFIG_FILE).st_mode & 0o777) != '0o600':
            issues.append(f"🟠 Incorrect permissions for config file - should be 600")
    
    # Check database encryption
    if not os.path.exists(DB_FILE):
        issues.append("🔴 Database not initialized")
    
    # Generate report
    if not issues:
        print("Overall Status: 🟢 SECURE")
        print("\nNo security issues detected")
    else:
        print("Overall Status: 🔴 VULNERABLE")
        print("\nSECURITY ISSUES DETECTED:")
        for issue in issues:
            print(f" - {issue}")
    
    print("\nSECURITY RECOMMENDATIONS:")
    print("  - Rotate encryption keys quarterly")
    print("  - Limit physical access to operation devices")
    print("  - Use secure networks for intelligence gathering")
    print("  - Perform regular security audits")
    print("="*50)

async def view_high_priority_alerts():
    """Display high-priority alerts from database"""
    try:
        async with aiosqlite.connect(DB_FILE) as db:
            cursor = await db.execute("""
            SELECT * FROM high_priority_alerts 
            ORDER BY timestamp DESC 
            LIMIT 20
            """)
            alerts = await cursor.fetchall()
            
            if not alerts:
                print("\n⚠️ No high-priority alerts recorded")
                return
                
            print("\n🚨 HIGH-PRIORITY ALERT HISTORY")
            print("=" * 120)
            print(f"{'ID':<4} | {'Threat':<6} | {'Conflict':<15} | {'Source':<20} | {'Timestamp':<20} | {'Title'}")
            print("-" * 120)
            for alert in alerts:
                alert_id = alert[0]
                threat_score = alert[4]
                conflict = alert[5]
                source = alert[3]
                timestamp = alert[7]
                title = alert[1][:70] + '...' if len(alert[1]) > 70 else alert[1]
                
                # Color coding based on threat level
                if threat_score >= 90:
                    color = "red"
                elif threat_score >= 80:
                    color = "yellow"
                else:
                    color = "white"
                    
                print(colored(f"{alert_id:<4} | {threat_score:<6} | {conflict:<15} | {source:<20} | {timestamp[:16]:<20} | {title}", color))
            print("=" * 120)
    except Exception as e:
        logger.error(f"Database error: {str(e)}")

def manage_configuration():
    """Manage system configuration settings"""
    global config
    
    print("\n⚙️ CONFIGURATION SETTINGS")
    print("=" * 50)
    for key, value in config['DEFAULT'].items():
        print(f"{key}: {value}")
    
    print("\n1. Change setting")
    print("2. Restore defaults")
    print("3. Security Recommendations")
    print("0. Back to main menu")
    choice = input("\n[CONFIG] Select option: ")
    
    if choice == '1':
        setting = input("Enter setting name: ")
        if setting in config['DEFAULT']:
            new_value = input(f"Enter new value for {setting}: ")
            config['DEFAULT'][setting] = new_value
            with open(CONFIG_FILE, 'w') as configfile:
                config.write(configfile)
            print(f"✅ Updated {setting} to {new_value}")
        else:
            print("❌ Invalid setting name")
            
    elif choice == '2':
        # Create backup of current config
        backup_file = os.path.join(CONFIG_DIR, f"mgcis_config_backup_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.ini")
        shutil.copy(CONFIG_FILE, backup_file)
        
        # Create default configuration
        config = configparser.ConfigParser()
        config['DEFAULT'] = {
            'encryption_key': base64.urlsafe_b64encode(os.urandom(32)).decode(),
            'threat_scan_interval': '3600',
            'alert_threshold': '80',
            'max_concurrent_requests': '15',
            'social_media_enabled': 'True',
            'geospatial_enabled': 'True',
            'llm_integration': 'False',
            'llm_api_key': '',
            'twitter_api_key': '',
            'twitter_api_secret': '',
            'twitter_access_token': '',
            'twitter_access_secret': '',
            'twitter_bearer_token': '',
            'audio_enabled': 'True',
            'audio_voice': 'f1',
            'audio_language': 'en',
            'version': '2.0.0'
        }
        
        with open(CONFIG_FILE, 'w') as configfile:
            config.write(configfile)
        print("✅ Restored default configuration")
        
    elif choice == '3':
        print("\n🔐 SECURITY RECOMMENDATIONS:")
        print("  - Rotate encryption keys quarterly")
        print("  - Set strong passwords for all system accounts")
        print("  - Limit physical access to operation devices")
        print("  - Use VPN for all intelligence operations")
        print("  - Regularly update the MG-CIS software")
        print("  - Perform security audits monthly")
        
    elif choice == '0':
        return
        
    else:
        print("❌ Invalid selection")

async def audio_briefing_menu(audio, intel_data):
    """Handle audio briefing menu options"""
    print("\n🎧 AUDIO BRIEFING OPTIONS")
    print("="*50)
    print("a. Critical Update (Top Alert)")
    print("b. Complete Audio Briefing (All Critical Alerts)")
    print("0. Back to Main Menu")
    choice = input("\nSelect audio briefing option: ")
    
    # Get critical alerts (score >= 80)
    critical_alerts = [a for a in intel_data if a['score'] >= 80] if intel_data else []
    
    if choice.lower() == 'a':
        if critical_alerts:
            print("🔊 Playing critical update...")
            audio.play_critical_update(critical_alerts)
        else:
            print("⚠️ No critical alerts available")
    elif choice.lower() == 'b':
        if critical_alerts:
            print("🔊 Playing complete briefing...")
            audio.play_complete_briefing(critical_alerts)
        else:
            print("⚠️ No critical alerts available")
    elif choice == '0':
        return
    else:
        print("❌ Invalid selection")

# ======================
# MAIN EXECUTION
# ======================
async def main():
    # Initialize system
    asyncio.set_event_loop_policy(asyncio.DefaultEventLoopPolicy())
    print(generate_ascii_banner())
    
    # Initialize database
    if not await init_db():
        print("❌ Critical error: Database initialization failed")
        return
    
    # Create intelligence engine
    async with IntelEngine() as intel_engine:
        # System state
        intel_data = []
        audio = AudioBriefing()
        
        while True:
            choice = display_main_menu()
            
            if choice == '1':  # Threat Scan
                print("\n🔍 Scanning global conflict sources...")
                start_time = time.time()
                with tqdm(total=100, bar_format="{l_bar}{bar}| {n_fmt}/{total_fmt}") as pbar:
                    intel_data = await intel_engine.analyze_global_threats(disable_audio=True)
                    pbar.update(100)
                elapsed = time.time() - start_time
                print(f"\n✅ Scan completed in {elapsed:.2f} seconds")
                display_threats(intel_data)
                
            elif choice == '2':  # Discover New Sources
                print("\n🕸️ Searching for new intelligence channels...")
                with tqdm(total=len(WEBSITE_LIST), desc="Discovering sources") as pbar:
                    new_feeds = await intel_engine.discover_feeds_from_website_list()
                    pbar.update(len(WEBSITE_LIST))
                print(f"\n✅ Discovered {len(new_feeds)} new intelligence sources")
                
            elif choice == '3':  # Daily Statistics
                print("\n📊 Calculating daily statistics...")
                stats = await intel_engine.get_daily_stats()
                display_daily_stats(stats)
                
            elif choice == '4':  # Conflict Statistics
                print("\n🔥 Gathering conflict statistics...")
                conflict_stats = await intel_engine.get_conflict_stats()
                display_conflict_stats(conflict_stats)
                
            elif choice == '5':  # Generate Report
                if not intel_data:
                    print("\n⚠️ No intelligence data. Run Threat Scan first.")
                else:
                    print("\n📊 Generating intelligence report...")
                    stats = await intel_engine.get_daily_stats()
                    report_data = await intel_engine.generate_detailed_report(intel_data, stats)
                    
                    print("\n📝 Select report format:")
                    print("1. JSON")
                    print("2. PDF")
                    print("3. DOCX")
                    print("4. Audio Briefing")
                    print("5. CSV")
                    print("6. Excel")
                    fmt_choice = input("\nChoose format (1-6): ")
                    
                    formats = {1: "json", 2: "pdf", 3: "docx", 4: "audio", 5: "csv", 6: "excel"}
                    fmt = formats.get(int(fmt_choice), "json")
                    
                    report_file = intel_engine.report_gen.generate_detailed_report(report_data, fmt)
                    if report_file:
                        print(f"✅ Report saved to: {report_file}")
                    else:
                        print("⚠️ Failed to generate report")
                    
            elif choice == '6':  # Geospatial Intelligence
                await view_geospatial_intelligence(intel_engine, intel_data)
                
            elif choice == '7':  # Threat Timeline
                await view_threat_timeline(intel_engine, intel_data)
                
            elif choice == '8':  # Social Media Monitoring
                print("\n📱 Monitoring social media...")
                social_data = await intel_engine.monitor_social_media()
                if social_data:
                    print(f"\n✅ Collected {len(social_data)} social media posts")
                    for post in social_data[:3]:
                        print(f"\n📱 {post['platform']} - @{post['author']}")
                        print(f"   {post['content']}")
                        print(f"   Sentiment: {post['sentiment']:.2f} | Conflict: {post['conflict_group']}")
                else:
                    print("⚠️ No social media data collected")
                
            elif choice == '9':  # Historical Conflicts
                await view_historical_conflicts()
                
            elif choice == '10':  # Create Custom Tool
                description = input("Enter tool description: ")
                tool_name, tool_path = await intel_engine.tool_manager.create_tool(description)
                print(f"✅ Created tool: {tool_name} at {tool_path}")
                
            elif choice == '11':  # Run Custom Tool
                await view_custom_tools(intel_engine.tool_manager)
                
            elif choice == '12':  # System Status
                print("\n⚙️ SYSTEM STATUS")
                print("="*50)
                print(f"OS: {platform.system()} {platform.release()} ({platform.machine()})")
                print(f"Python: {platform.python_version()}")
                print(f"Database: {DB_FILE}")
                print(f"Articles: {len(intel_data)} in memory")
                print(f"Config Version: {config['DEFAULT'].get('version', '2.0.0')}")
                print("Status: OPERATIONAL")
                print("="*50)
                
            elif choice == '13':  # Security Audit
                security_audit()
                
            elif choice == '14':  # View High-Priority Alerts
                await view_high_priority_alerts()
                
            elif choice == '15':  # Configuration Settings
                manage_configuration()
                
            elif choice == '16':  # Audio Voice Selection
                audio.select_voice()
                
            elif choice == '17':  # Live Dashboard
                intel_engine.dashboard.start(intel_data)
                
            elif choice == '18':  # Audio Briefings
                if intel_data:
                    await audio_briefing_menu(audio, intel_data)
                else:
                    print("\n⚠️ No intelligence data. Run Threat Scan first.")
                
            elif choice == '0':  # Exit
                intel_engine.dashboard.stop()
                print("\n✅ Operation complete. Stay vigilant.")
                break
                
            else:
                print("❌ Invalid selection. Please choose 0-18.")

if __name__ == "__main__":
    try:
        asyncio.run(main())
    except KeyboardInterrupt:
        print("\n🛑 Operation interrupted. Connection Terminated...")
        sys.exit(0)